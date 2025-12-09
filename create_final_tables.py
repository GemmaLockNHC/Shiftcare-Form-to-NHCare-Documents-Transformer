#!/usr/bin/env python3

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import csv
import os
import re
import io
import tempfile

# Font registration - lazy loaded to avoid slow startup
_VERDANA_FONT = None
_CALIBRI_FONT = None
_CALIBRI_BOLD_FONT = None
_FONTS_REGISTERED = False

def _register_fonts():
    """Lazy-load font registration - only called when needed"""
    global _VERDANA_FONT, _CALIBRI_FONT, _CALIBRI_BOLD_FONT, _FONTS_REGISTERED
    
    if _FONTS_REGISTERED:
        return  # Already registered
    
    _VERDANA_FONT = 'Helvetica'  # Default fallback
    _CALIBRI_FONT = 'Helvetica'  # Default fallback
    _CALIBRI_BOLD_FONT = 'Helvetica-Bold'
    
    try:
        import platform
        system = platform.system()
        
        # Only check paths relevant to current platform
        if system == 'Darwin':  # macOS
            verdana_paths = ['/System/Library/Fonts/Supplemental/Verdana.ttf']
            calibri_paths = [
                '/System/Library/Fonts/Supplemental/Calibri.ttf',
                '/Library/Fonts/Calibri.ttf',
            ]
            calibri_bold_paths = [
                '/System/Library/Fonts/Supplemental/Calibri Bold.ttf',
                '/Library/Fonts/Calibri Bold.ttf',
            ]
            home = os.path.expanduser('~')
            calibri_paths.append(os.path.join(home, 'Library/Fonts/Calibri.ttf'))
            calibri_bold_paths.append(os.path.join(home, 'Library/Fonts/Calibri Bold.ttf'))
        elif system == 'Windows':
            verdana_paths = ['C:/Windows/Fonts/verdana.ttf']
            calibri_paths = ['C:/Windows/Fonts/calibri.ttf', 'C:/Windows/Fonts/CALIBRI.TTF']
            calibri_bold_paths = ['C:/Windows/Fonts/calibrib.ttf', 'C:/Windows/Fonts/CALIBRIB.TTF']
        else:  # Linux (including Render)
            verdana_paths = ['/usr/share/fonts/truetype/ttf-dejavu/DejaVuSans.ttf']
            calibri_paths = [
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                '/usr/share/fonts/truetype/msttcorefonts/arial.ttf',
            ]
            calibri_bold_paths = [
                '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
                '/usr/share/fonts/truetype/msttcorefonts/arialbd.ttf',
            ]
        
        # Register Verdana
        for path in verdana_paths:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont('Verdana', path))
                    _VERDANA_FONT = 'Verdana'
                    print(f"Verdana font registered from: {path}")
                    break
                except Exception:
                    continue
        
        # Register Calibri
        calibri_registered = False
        calibri_bold_registered = False
        for path in calibri_paths:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont('Calibri', path))
                    calibri_registered = True
                    _CALIBRI_FONT = 'Calibri'
                    print(f"Calibri font registered from: {path}")
                    break
                except Exception:
                    continue
        
        for path in calibri_bold_paths:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont('Calibri-Bold', path))
                    calibri_bold_registered = True
                    _CALIBRI_BOLD_FONT = 'Calibri-Bold'
                    print(f"Calibri Bold font registered from: {path}")
                    break
                except Exception:
                    continue
        
        if not calibri_registered:
            print("Calibri font not found, will use Helvetica as fallback")
        if not _VERDANA_FONT == 'Verdana':
            print("Verdana font not found, will use Helvetica as fallback")
    except Exception as e:
        print(f"Could not register fonts: {e}, using Helvetica")
    
    _FONTS_REGISTERED = True

# Font getters - call _register_fonts() when first accessed
def get_verdana_font():
    if not _FONTS_REGISTERED:
        _register_fonts()
    return _VERDANA_FONT or 'Helvetica'

def get_calibri_font():
    if not _FONTS_REGISTERED:
        _register_fonts()
    return _CALIBRI_FONT or 'Helvetica'

def get_calibri_bold_font():
    if not _FONTS_REGISTERED:
        _register_fonts()
    return _CALIBRI_BOLD_FONT or 'Helvetica-Bold'

# For backward compatibility - these will be set after first call
# Use get_verdana_font(), get_calibri_font(), get_calibri_bold_font() instead
VERDANA_FONT = 'Helvetica'  # Will be updated on first use
CALIBRI_FONT = 'Helvetica'  # Will be updated on first use
CALIBRI_BOLD_FONT = 'Helvetica-Bold'  # Will be updated on first use

# Try to import Excel library for formatted output
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    openpyxl_available = True
except Exception:
    openpyxl_available = False

# Try to import PDF parsing libraries
try:
    import pdfplumber  # text extraction
except Exception:
    pdfplumber = None
try:
    from pypdf import PdfReader  # form fields
except Exception:
    PdfReader = None

# Define custom colors
BLUE_COLOR = colors.HexColor('#316DB2')

def load_ndis_support_items():
    """Load NDIS support items from CSV file and return as a dictionary for lookup"""
    ndis_items = {}
    try:
        with open('outputs/other/NDIS Support Items - NDIS Support Items.csv', 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            for row in reader:
                # Use support item name as key for lookup
                item_name = row['Support Item Name'].strip()
                ndis_items[item_name] = {
                    'number': row['Support Item Number'].strip(),
                    'unit': row['Unit'].strip(),
                    'wa_price': row.get('WA', '').strip(),
                    'qld_price': row.get('QLD', '').strip()
                }
        print(f"DEBUG: Loaded {len(ndis_items)} NDIS support items from CSV")
        # Verify establishment fee item exists
        if "Establishment Fee For Personal Care/Participation" in ndis_items:
            est_fee = ndis_items["Establishment Fee For Personal Care/Participation"]
            print(f"DEBUG: Establishment fee item found - WA: {est_fee.get('wa_price')}, QLD: {est_fee.get('qld_price')}")
        else:
            print("DEBUG: WARNING - Establishment Fee For Personal Care/Participation not found in NDIS items")
    except FileNotFoundError:
        print("NDIS Support Items CSV file not found. Using placeholder data.")
    except Exception as e:
        print(f"Error loading NDIS support items: {e}")
    
    return ndis_items

def get_price_state(team_value):
    """
    Determine which state's prices to use based on the team.
    
    Args:
        team_value: The team name from 'Neighbourhood Care representative team'
    
    Returns:
        str: 'WA' or 'QLD'
    """
    if not team_value:
        return 'WA'  # Default to WA if team is not specified
    
    team_value_clean = team_value.strip().lower()
    
    # WA teams
    wa_teams = ['fremantle', 'belmont', 'metro-x', 'rockingham', 'wanneroo']
    if team_value_clean in wa_teams:
        return 'WA'
    
    # QLD teams
    qld_teams = ['brisbane', 'beaudesert', 'ipswich', 'gold coast']
    if team_value_clean in qld_teams:
        return 'QLD'
    
    # Default to WA if team doesn't match
    return 'WA'

def lookup_support_item(ndis_items, item_name):
    """Look up a support item by name and return its details"""
    if item_name in ndis_items:
        return ndis_items[item_name]
    else:
        # Try partial matching for common support items
        for key, value in ndis_items.items():
            if item_name.lower() in key.lower() or key.lower() in item_name.lower():
                return value
        # Return placeholder if not found
        return {
            'number': '[Not Found]',
            'unit': 'Hour',
            'wa_price': '$0.00',
            'qld_price': '$0.00'
        }

def get_establishment_fee(csv_data, ndis_items, team_value=None):
    """
    Calculate the establishment fee based on client status and support hours.
    
    The establishment fee is applied only if:
    - Client is new (isNewClient == "Yes")
    - Client is receiving 20+ hours of support (isReceiving20HoursSupport == "Yes" or similar)
    
    Args:
        csv_data: Dictionary containing form data
        ndis_items: Dictionary of NDIS support items loaded from CSV
        team_value: The team name from 'Neighbourhood Care representative team' (optional)
    
    Returns:
        str: Formatted establishment fee amount (e.g., "$702.30" or "$0.00")
    """
    # Check for new client status - try various possible field names
    # JavaScript code checked: submission.isNewClient == "Yes"
    is_new_client = False
    new_client_fields = [
        'isNewClient',  # Primary field name from JavaScript
        'Is this client new to Neighbourhood Care?',  # Actual PDF question text
        'Is this a new client?',
        'Is the client new?',
        'Is this client new?',
        'New client',
        'Is new client',
        'Client is new'
    ]
    
    for field in new_client_fields:
        value = csv_data.get(field, '').strip()
        # Clean up checkbox characters and other special characters
        value = value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
        # Match JavaScript: exact "Yes" check (case-sensitive in JS, but we'll be flexible)
        if value == "Yes" or normalize_key(value) == 'yes':
            is_new_client = True
            break
    
    # Check for 20+ hours of support - try various possible field names
    # JavaScript code checked: submission.isReceiving20HoursSupport == "Yes"
    # Also had: submission.isReceiving20HoursOfSupport == "Yes"
    is_receiving_20_hours = False
    hours_support_fields = [
        'isReceiving20HoursSupport',  # Primary field name from JavaScript
        'isReceiving20HoursOfSupport',  # Alternative from JavaScript
        'Is Neighbourhood Care delivering 20 or more hours of support per month?',  # Actual PDF question text
        'Is the client receiving 20 or more hours of support?',
        'Is receiving 20 hours of support',
        'Is receiving 20+ hours of support',
        'Receiving 20 hours support',
        'Receiving 20+ hours',
        '20 hours support',
        '20 or more hours of support'
    ]
    
    for field in hours_support_fields:
        value = csv_data.get(field, '').strip()
        # Clean up checkbox characters and other special characters
        value = value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
        # Match JavaScript: exact "Yes" check (case-sensitive in JS, but we'll be flexible)
        if value == "Yes" or normalize_key(value) == 'yes':
            is_receiving_20_hours = True
            break
    
    # If both conditions are met, get the establishment fee from NDIS support items
    if is_new_client and is_receiving_20_hours:
        # Look up "Establishment Fee For Personal Care/Participation" in NDIS items
        establishment_fee_item = lookup_support_item(ndis_items, "Establishment Fee For Personal Care/Participation")
        
        # Determine which state's price to use
        price_state = get_price_state(team_value)
        price_key = 'wa_price' if price_state == 'WA' else 'qld_price'
        price = establishment_fee_item.get(price_key, '$0.00')
        
        # Clean up the price string (remove any extra formatting)
        price = price.strip()
        
        # If price is empty or invalid, try to get it from the item directly
        if not price or price == '' or price == '$0.00':
            # Try the other price key as fallback
            other_key = 'qld_price' if price_key == 'wa_price' else 'wa_price'
            price = establishment_fee_item.get(other_key, '$0.00').strip()
        
        # Ensure price has $ prefix and proper formatting
        if price and price != '$0.00':
            if not price.startswith('$'):
                # If it's a number, add $ prefix
                try:
                    # Remove any commas and $ signs, then format
                    price_num = float(price.replace('$', '').replace(',', ''))
                    price = f"${price_num:.2f}"
                except (ValueError, AttributeError):
                    price = '$0.00'
            else:
                # Price already has $, just ensure proper formatting
                try:
                    price_num = float(price.replace('$', '').replace(',', ''))
                    price = f"${price_num:.2f}"
                except (ValueError, AttributeError):
                    price = '$0.00'
        
        # Debug output
        if price == '$0.00':
            print(f"DEBUG: Establishment fee calculation - is_new_client: {is_new_client}, is_receiving_20_hours: {is_receiving_20_hours}")
            print(f"DEBUG: Price state: {price_state}, Price key: {price_key}")
            print(f"DEBUG: Establishment fee item found: {establishment_fee_item}")
        
        return price
    else:
        # Debug output when conditions not met
        print(f"DEBUG: Establishment fee conditions not met - is_new_client: {is_new_client}, is_receiving_20_hours: {is_receiving_20_hours}")
        return '$0.00'

def load_active_users(team_value=None):
    """
    Load active users from CSV file and return as a dictionary for lookup.
    
    Args:
        team_value: The team name to determine which CSV file to use.
                   QLD teams (Beaudesert, Brisbane, Gold Coast, Ipswich) use Active_Users_1763520740.csv
                   Other teams use Active_Users_1761707021.csv
    
    Returns:
        Dictionary of active users keyed by name
    """
    active_users = {}
    
    # Determine which CSV file to use based on team
    qld_teams = ['beaudesert', 'brisbane', 'gold coast', 'ipswich']
    team_lower = team_value.strip().lower() if team_value else ''
    
    if team_lower in qld_teams:
        csv_filename = 'outputs/other/Active_Users_1763520740.csv'
        print(f"DEBUG: Using QLD active users CSV for team: {team_value}")
    else:
        csv_filename = 'outputs/other/Active_Users_1761707021.csv'
        print(f"DEBUG: Using default active users CSV for team: {team_value or 'unknown'}")
    
    try:
        with open(csv_filename, 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            for row in reader:
                # Use name as key for lookup
                user_name = row['name'].strip()
                active_users[user_name] = {
                    'name': row['name'].strip(),
                    'mobile': row['mobile'].strip(),
                    'email': row['email'].strip(),
                    'team': (row.get('area') or row.get('role') or '').strip()
                }
        print(f"DEBUG: Loaded {len(active_users)} active users from {csv_filename}")
    except FileNotFoundError:
        print(f"Active Users CSV file not found: {csv_filename}. Using placeholder data.")
    except Exception as e:
        print(f"Error loading active users from {csv_filename}: {e}")
    
    return active_users

def lookup_user_data(active_users, respondent_name):
    """Look up user data by respondent name and return contact details"""
    if respondent_name in active_users:
        return active_users[respondent_name]
    else:
        # Try partial matching
        for key, value in active_users.items():
            if respondent_name.lower() in key.lower() or key.lower() in respondent_name.lower():
                return value
        # Return placeholder if not found
        return {
            'name': respondent_name,
            'mobile': '[Not Found]',
            'email': '[Not Found]'
        }

def normalize_key(key: str) -> str:
    """Normalize a key for comparison"""
    return str(key or "").strip().lower()

def extract_pdf_fields_pdfreader(pdf_path: str) -> dict:
    """Extract form fields from PDF using PdfReader"""
    if PdfReader is None:
        return {}
    try:
        reader = PdfReader(pdf_path)
        fields = {}
        if reader.get_fields():
            for name, field in reader.get_fields().items():
                value = field.get("/V")
                if value is None:
                    continue
                fields[name] = str(value)
        return fields
    except Exception:
        return {}

def extract_pdf_text_pdfplumber(pdf_path: str) -> str:
    """Extract all text from PDF using pdfplumber"""
    if pdfplumber is None:
        return ""
    try:
        text_parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except Exception:
        return ""

def parse_pdf_to_data(pdf_path: str) -> dict:
    """Parse PDF and extract data, mapping to CSV field names"""
    data = {}
    
    # Debug flag - set to True to see what's being extracted
    DEBUG = False
    
    # First try to extract form fields
    fields = extract_pdf_fields_pdfreader(pdf_path)
    if fields:
        # Debug: print all field names found (commented out for production)
        # print("PDF Form Fields found:", list(fields.keys()))
        
        # Map form fields to CSV field names
        # The PDF form field NAMES are simple like "First name", "Home address", etc.
        # The VALUES in those fields are the actual data (might look like "First name (Details of the Client)" but treat as real data)
        def find_in_fields(*candidates):
            # Try all candidates - match simple field names
            for cand in candidates:
                cand_norm = normalize_key(cand)
                for key, val in fields.items():
                    key_norm = normalize_key(key)
                    value = str(val).strip()
                    if not value:
                        continue
                    
                    # Exact match
                    if cand_norm == key_norm:
                        return value
                    
                    # Substring match - candidate is in the key
                    if cand_norm in key_norm:
                        return value
            return ""
        
        # Details of the Client section - match simple field names like "First name"
        data['First name (Details of the Client)'] = find_in_fields("first name", "firstname")
        data['Middle name (Details of the Client)'] = find_in_fields("middle name", "middlename")
        data['Surname (Details of the Client)'] = find_in_fields("surname", "family name", "last name", "lastname")
        data['NDIS number (Details of the Client)'] = find_in_fields("ndis number", "ndis")
        data['Date of birth (Details of the Client)'] = find_in_fields("date of birth", "dob", "birth date")
        data['Gender (Details of the Client)'] = find_in_fields("gender")
        
        # Contact Details of the Client section - match simple field names like "Home address"
        data['Home address (Contact Details of the Client)'] = find_in_fields("home address", "address")
        data['Home phone (Contact Details of the Client)'] = find_in_fields("home phone", "homephone")
        data['Work phone (Contact Details of the Client)'] = find_in_fields("work phone", "workphone")
        data['Mobile phone (Contact Details of the Client)'] = find_in_fields("mobile phone", "mobile", "mobilephone")
        data['Email address (Contact Details of the Client)'] = find_in_fields("email address", "email")
        
        # Emergency contact - match simple field names, but need to distinguish from client fields
        # Try emergency-specific names first, then fallback to simple names
        data['First name (Emergency contact)'] = find_in_fields("first name (emergency contact)", "emergency contact first name", "emergency first name")
        if not data['First name (Emergency contact)']:
            # If we found a "first name" field, check if there's an emergency-specific one
            # For now, we'll rely on text extraction to get the right one
            pass
        data['Surname (Emergency contact)'] = find_in_fields("surname (emergency contact)", "emergency contact surname", "emergency contact last name", "emergency surname", "emergency last name")
        data['Is the primary carer also the emergency contact for the participant?'] = find_in_fields("primary carer also emergency contact", "is primary carer emergency contact")
        data['Home phone (Emergency contact)'] = find_in_fields("home phone (emergency contact)", "emergency contact home phone", "emergency home phone")
        data['Mobile phone (Emergency contact)'] = find_in_fields("mobile phone (emergency contact)", "emergency contact mobile phone", "emergency mobile phone", "emergency contact mobile")
        data['Work phone (Emergency contact)'] = find_in_fields("work phone (emergency contact)", "emergency contact work phone", "emergency work phone")
        # ONLY look for "relationship to client" field - must be specifically under Emergency contact
        # Don't do fuzzy matching - only exact match for emergency contact relationship
        data['Relationship to client (Emergency contact)'] = find_in_fields("relationship to client (emergency contact)", "relationship to client")
        
        # Extract Person Signing the Agreement fields
        person_signing = find_in_fields("person signing the agreement", "who is signing", "signatory")
        # Clean up checkbox characters
        if person_signing:
            person_signing = person_signing.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
        data['Person signing the agreement'] = person_signing
        data['First name (Person Signing the Agreement)'] = find_in_fields("first name (person signing the agreement)", "first name (person signing", "person signing first name", "signatory first name")
        data['Surname (Person Signing the Agreement)'] = find_in_fields("surname (person signing the agreement)", "surname (person signing", "person signing surname", "person signing last name", "signatory surname", "signatory last name")
        data['Relationship to client (Person Signing the Agreement)'] = find_in_fields("relationship to client (person signing the agreement)", "relationship to client (person signing", "person signing relationship", "signatory relationship")
        data['Home address (Person Signing the Agreement)'] = find_in_fields("home address (person signing the agreement)", "home address (person signing", "person signing address", "signatory address")
        data['Home phone (Person Signing the Agreement)'] = find_in_fields("home phone (person signing the agreement)", "home phone (person signing", "person signing home phone", "signatory home phone")
        data['Mobile phone (Person Signing the Agreement)'] = find_in_fields("mobile phone (person signing the agreement)", "mobile phone (person signing", "person signing mobile", "signatory mobile")
        data['Email address (Person Signing the Agreement)'] = find_in_fields("email address (person signing the agreement)", "email address (person signing", "person signing email", "signatory email")
        
        # Extract Primary carer fields
        data['First name (Primary carer)'] = find_in_fields("first name (primary carer)", "first name (primary carer", "primary carer first name")
        data['Surname (Primary carer)'] = find_in_fields("surname (primary carer)", "surname (primary carer", "primary carer surname", "primary carer last name")
        data['Relationship to client (Primary carer)'] = find_in_fields("relationship to client (primary carer)", "relationship to client (primary carer", "primary carer relationship")
        data['Home address (Primary carer)'] = find_in_fields("home address (primary carer)", "home address (primary carer", "primary carer address")
        data['Home phone (Primary carer)'] = find_in_fields("home phone (primary carer)", "home phone (primary carer", "primary carer home phone")
        data['Mobile phone (Primary carer)'] = find_in_fields("mobile phone (primary carer)", "mobile phone (primary carer", "primary carer mobile")
        data['Email address (Primary carer)'] = find_in_fields("email address (primary carer)", "email address (primary carer", "primary carer email")
    
    # Always try text extraction as well to fill in any missing fields
    # This ensures we get all fields even if form field extraction missed some
    text = extract_pdf_text_pdfplumber(pdf_path)
    if text:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        
        # Identify section boundaries
        section_starts = []
        for i, line in enumerate(lines):
            line_lower = normalize_key(line)
            line_clean = line.strip()
            
            # Only match actual section headers - they should be standalone lines without parentheses
            # Section headers are typically short and don't contain field values
            is_section_header = (
                len(line_clean) < 50 and  # Section headers are short
                '(' not in line_clean and  # Section headers don't have parentheses (values do)
                ':' not in line_clean      # Section headers don't have colons
            )
            
            if is_section_header:
                if "details of the client" in line_lower and "contact" not in line_lower:
                    section_starts.append(("details", i))
                elif "contact details of the client" in line_lower:
                    section_starts.append(("contact", i))
                elif "primary carer" in line_lower and "emergency" not in line_lower:
                    section_starts.append(("primary_carer", i))
                elif "emergency contact" in line_lower:
                    section_starts.append(("emergency", i))
                elif any(x in line_lower for x in ["needs of the client", "ndis information", "support items", "formal supports", 
                                                   "important people", "home life", "health information", 
                                           "care requirements", "behaviour requirements", "other information", "consents"]):
                    # End of relevant sections - but only break if we've found primary_carer and emergency sections
                    if any(sec[0] == "primary_carer" for sec in section_starts) and any(sec[0] == "emergency" for sec in section_starts):
                        break
        
        # Helper function to find value in a specific section - SIMPLIFIED
        def find_value_in_section(label_patterns, section_type):
            """Find value only in the specified section - just get the next line after the label"""
            # Find the relevant section
            section_start = None
            section_end = None
            
            for sec_type, start_idx in section_starts:
                if sec_type == section_type:
                    section_start = start_idx
                    # Find end of this section (start of next section or end of lines)
                    for next_sec_type, next_start_idx in section_starts:
                        if next_start_idx > start_idx:
                            section_end = next_start_idx
                            break
                    if section_end is None:
                        section_end = len(lines) if lines else 0
                    break
            
            if section_start is None or section_end is None:
                return ""
            
            # Ensure both are integers
            section_start = int(section_start) if section_start is not None else 0
            section_end = int(section_end) if section_end is not None else len(lines) if lines else 0
            
            # Only search within this section
            for i in range(section_start, section_end):
                line = lines[i]
                line_lower = normalize_key(line)
                
                for pattern in label_patterns:
                    pattern_lower = normalize_key(pattern)
                    
                    # Simple match - check if pattern matches the line
                    line_clean = line_lower.replace("(details of the client)", "").replace("(contact details of the client)", "").strip()
                    pattern_clean = pattern_lower.replace("(details of the client)", "").replace("(contact details of the client)", "").strip()
                    
                    matches = (
                        pattern_lower == line_lower or
                        pattern_clean == line_clean or
                        line_lower.startswith(pattern_lower)
                    )
                    
                    if matches:
                        # Look for value on same line after colon (if present)
                        if ':' in line:
                            parts = line.split(':', 1)
                            if len(parts) > 1 and parts[1].strip():
                                return parts[1].strip()
                        
                        # Just get the next non-empty line - that's the value
                        # But skip if it's clearly another field label
                        field_labels = ['first name', 'middle name', 'surname', 'ndis number', 'date of birth', 'gender',
                                       'home address', 'home phone', 'work phone', 'mobile phone', 'email address',
                                       'preferred name', 'key code', 'postal address', 'preferred method of contact',
                                       'relationship to client']
                        section_end_safe = section_end if section_end is not None and isinstance(section_end, int) else (len(lines) if lines else i + 5)
                        for j in range(i + 1, min(i + 5, section_end_safe)):
                            next_line = lines[j].strip()
                            if not next_line or next_line in ['•', '●', '○', '☐', '☑', '✓']:
                                continue
                            
                            # Skip if it's another field label
                            next_line_lower = normalize_key(next_line)
                            is_field_label = False
                            for fl in field_labels:
                                if fl == next_line_lower or (fl in next_line_lower and len(next_line) < 50 and '(' not in next_line):
                                    is_field_label = True
                                    break
                            
                            # Skip instruction text
                            if len(next_line) > 80 or any(x in next_line_lower for x in ['write', 'below', 'same as', 'if their']):
                                continue
                            
                            if not is_field_label:
                                return next_line
            
            return ""
        
        # Helper function for fields that aren't in specific sections - SIMPLIFIED
        def find_value_after_label(label_patterns, start_idx=0):
            for i in range(start_idx, len(lines) if lines else 0):
                line_lower = normalize_key(lines[i])
                for pattern in label_patterns:
                    pattern_lower = normalize_key(pattern)
                    # Match if pattern is in the line (but not if the line IS the pattern - that's the label)
                    if pattern_lower in line_lower:
                        # Look for value on same line after colon
                        if ':' in lines[i]:
                            parts = lines[i].split(':', 1)
                            if len(parts) > 1 and parts[1].strip():
                                return parts[1].strip()
                        # Skip the label line itself and get the next non-empty line - that's the value
                        for j in range(i + 1, min(i + 3, len(lines) if lines else i + 3)):
                            next_line = lines[j].strip()
                            if next_line and next_line not in ['•', '●', '○', '☐', '☑', '✓']:
                                # Make sure we're not returning the label itself
                                next_line_lower = normalize_key(next_line)
                                if next_line_lower != pattern_lower:
                                    return next_line
            return ""
    
        # Extract data using section-aware text parsing - only fill in missing fields
        if not data.get('First name (Details of the Client)'):
            data['First name (Details of the Client)'] = find_value_in_section(['First name', 'First name (Details of the Client)'], "details")
        if not data.get('Middle name (Details of the Client)'):
            data['Middle name (Details of the Client)'] = find_value_in_section(['Middle name', 'Middle name (Details of the Client)'], "details")
        if not data.get('Surname (Details of the Client)'):
            data['Surname (Details of the Client)'] = find_value_in_section(['Surname', 'Surname (Details of the Client)', 'Family name', 'Last name'], "details")
        if not data.get('NDIS number (Details of the Client)'):
            data['NDIS number (Details of the Client)'] = find_value_in_section(['NDIS number', 'NDIS number (Details of the Client)'], "details")
        if not data.get('Date of birth (Details of the Client)'):
            data['Date of birth (Details of the Client)'] = find_value_in_section(['Date of birth', 'Date of birth (Details of the Client)', 'DOB'], "details")
        if not data.get('Gender (Details of the Client)'):
            data['Gender (Details of the Client)'] = find_value_in_section(['Gender', 'Gender (Details of the Client)'], "details")
        if not data.get('Home address (Contact Details of the Client)'):
            data['Home address (Contact Details of the Client)'] = find_value_in_section(['Home address', 'Home address (Contact Details of the Client)', 'Address'], "contact")
        if not data.get('Home phone (Contact Details of the Client)'):
            data['Home phone (Contact Details of the Client)'] = find_value_in_section(['Home phone', 'Home phone (Contact Details of the Client)'], "contact")
        if not data.get('Work phone (Contact Details of the Client)'):
            data['Work phone (Contact Details of the Client)'] = find_value_in_section(['Work phone', 'Work phone (Contact Details of the Client)'], "contact")
        if not data.get('Mobile phone (Contact Details of the Client)'):
            data['Mobile phone (Contact Details of the Client)'] = find_value_in_section(['Mobile phone', 'Mobile phone (Contact Details of the Client)'], "contact")
        if not data.get('Email address (Contact Details of the Client)'):
            data['Email address (Contact Details of the Client)'] = find_value_in_section(['Email address', 'Email address (Contact Details of the Client)', 'Email'], "contact")
    
        # Extract emergency contact fields - try emergency section first, then fallback to general search
        if not data.get('First name (Emergency contact)'):
            emergency_first = find_value_in_section(['First name'], "emergency")
            if emergency_first:
                data['First name (Emergency contact)'] = emergency_first
            else:
                data['First name (Emergency contact)'] = find_value_after_label(['First name (Emergency contact)'])
        if not data.get('Surname (Emergency contact)'):
            emergency_surname = find_value_in_section(['Surname'], "emergency")
            if emergency_surname:
                data['Surname (Emergency contact)'] = emergency_surname
            else:
                data['Surname (Emergency contact)'] = find_value_after_label(['Surname (Emergency contact)'])
        if not data.get('Is the primary carer also the emergency contact for the participant?'):
            data['Is the primary carer also the emergency contact for the participant?'] = find_value_after_label(['Is the primary carer also the emergency contact'])
        # Extract emergency contact phone and relationship fields
        if not data.get('Home phone (Emergency contact)'):
            data['Home phone (Emergency contact)'] = find_value_in_section(['Home phone'], "emergency")
        if not data.get('Mobile phone (Emergency contact)'):
            data['Mobile phone (Emergency contact)'] = find_value_in_section(['Mobile phone'], "emergency")
        if not data.get('Work phone (Emergency contact)'):
            data['Work phone (Emergency contact)'] = find_value_in_section(['Work phone'], "emergency")
        if not data.get('Relationship to client (Emergency contact)'):
            # ONLY search in the emergency section for "Relationship to client"
            # If not found, leave it empty - no fallback searches
            relationship = find_value_in_section(['Relationship to client'], "emergency")
            if relationship:
                data['Relationship to client (Emergency contact)'] = relationship
            # If not found, leave it empty - don't do fallback searches
        
        # Extract other fields that might be in the PDF
        if not data.get('Preferred method of contact'):
            data['Preferred method of contact'] = find_value_after_label(['Preferred method of contact', 'Preferred contact method'])
        if not data.get('Total core budget to allocate to Neighbourhood Care'):
            data['Total core budget to allocate to Neighbourhood Care'] = find_value_after_label(['Total core budget', 'core budget'])
        if not data.get('Total capacity building budget to allocate to Neighbourhood Care'):
            data['Total capacity building budget to allocate to Neighbourhood Care'] = find_value_after_label(['Total capacity building budget', 'capacity building budget'])
        if not data.get('Plan start date'):
            data['Plan start date'] = find_value_after_label(['Plan start date', 'Plan start'])
        if not data.get('Plan end date'):
            data['Plan end date'] = find_value_after_label(['Plan end date', 'Plan end'])
        if not data.get('Service start date'):
            data['Service start date'] = find_value_after_label(['Service start date', 'Service start'])
        if not data.get('Service end date'):
            data['Service end date'] = find_value_after_label(['Service end date', 'Service end'])
        # Always try text extraction for Person signing the agreement (form fields might return the label)
        person_signing_text = find_value_after_label(['Person signing the agreement', 'Who is signing'])
        if person_signing_text and person_signing_text.lower() != 'person signing the agreement':
            # Clean up checkbox characters
            person_signing_text = person_signing_text.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
            if person_signing_text:
                data['Person signing the agreement'] = person_signing_text
        if not data.get('First name (Person Signing the Agreement)'):
            data['First name (Person Signing the Agreement)'] = find_value_after_label(['First name (Person Signing the Agreement)'])
        if not data.get('Surname (Person Signing the Agreement)'):
            data['Surname (Person Signing the Agreement)'] = find_value_after_label(['Surname (Person Signing the Agreement)'])
        if not data.get('Relationship to client (Person Signing the Agreement)'):
            data['Relationship to client (Person Signing the Agreement)'] = find_value_after_label(['Relationship to client (Person Signing the Agreement)', 'Relationship'])
        if not data.get('Home address (Person Signing the Agreement)'):
            data['Home address (Person Signing the Agreement)'] = find_value_after_label(['Home address (Person Signing the Agreement)'])
        if not data.get('First name (Primary carer)'):
            # Look for "First name" label in Primary carer section
            data['First name (Primary carer)'] = find_value_in_section(['First name'], "primary_carer")
        if not data.get('Surname (Primary carer)'):
            # Look for "Surname" label in Primary carer section  
            data['Surname (Primary carer)'] = find_value_in_section(['Surname'], "primary_carer")
        if not data.get('Relationship to client (Primary carer)'):
            data['Relationship to client (Primary carer)'] = find_value_after_label(['Relationship to client (Primary carer)'])
        if not data.get('Home address (Primary carer)'):
            data['Home address (Primary carer)'] = find_value_after_label(['Home address (Primary carer)'])
        if not data.get('Plan management type'):
            data['Plan management type'] = find_value_after_label(['Plan management type', 'Plan management'])
        if not data.get('Plan manager name'):
            data['Plan manager name'] = find_value_after_label(['Plan manager name'])
        
        # Extract support items from Support Items Required section
        for i in range(1, 20):
            key = f'Support item ({i}) (Support Items Required)'
            if not data.get(key):
                # Look for "Support item (X)" label and get the value
                label_pattern = f'Support item ({i})'
                value = find_value_after_label([label_pattern])
                if value:
                    data[key] = value
        if not data.get('Plan manager postal address'):
            data['Plan manager postal address'] = find_value_after_label(['Plan manager postal address', 'Plan manager address'])
        if not data.get('Plan manager phone number'):
            data['Plan manager phone number'] = find_value_after_label(['Plan manager phone', 'Plan manager phone number'])
        if not data.get('Plan manager email address'):
            data['Plan manager email address'] = find_value_after_label(['Plan manager email'])
        if not data.get('Respondent'):
            data['Respondent'] = find_value_after_label(['Respondent', 'Neighbourhood Care representative'])
        if not data.get('Neighbourhood Care representative team'):
            data['Neighbourhood Care representative team'] = find_value_after_label(['Neighbourhood Care representative team', 'Team'])
        
        # Extract establishment fee related fields
        if not data.get('Is this client new to Neighbourhood Care?'):
            data['Is this client new to Neighbourhood Care?'] = find_value_after_label(['Is this client new to Neighbourhood Care?', 'Is this client new', 'Is this a new client'])
        if not data.get('Is Neighbourhood Care delivering 20 or more hours of support per month?'):
            data['Is Neighbourhood Care delivering 20 or more hours of support per month?'] = find_value_after_label(['Is Neighbourhood Care delivering 20 or more hours of support per month?', 'Is Neighbourhood Care delivering 20', '20 or more hours of support'])
    
        # Extract consent responses - look for Yes/No patterns
        consent_labels = [
            'I agree to receive services from Neighbourhood Care.',
            'I consent for Neighbourhood Care to create an NDIS portal service booking',
            'I understand that if at any time I (The Participant) require emergency medical assistance',
            'I agree that Neighbourhood Care staff may administer simple first aid',
            'I consent for Neighbourhood Care to discuss relevant information',
            'I agree not to smoke inside the home',
            'I understand that an Emergency Response Plan will be developed',
            'I consent for Neighbourhood Care for I (The Participant) to be photographed',
            'I give authority for my details or information to be shared'
        ]
        
        for consent_label in consent_labels:
            if consent_label not in data:
                # Look for the consent text and find Yes/No after it
                for i, line in enumerate(lines):
                    if normalize_key(consent_label.split('.')[0]) in normalize_key(line):
                        # Look for Yes/No in nearby lines
                        for j in range(max(0, i-2), min(len(lines) if lines else i+5, i+5)):
                            if normalize_key(lines[j]) in ['yes', 'no']:
                                data[consent_label] = lines[j]
                        break
    
        # Debug output
        if DEBUG:
            print("\n=== Extracted Data ===")
            for key, value in data.items():
                if value:
                    print(f"{key}: {value}")
            print("=====================\n")
    
    return data

def get_preferred_contact_details(csv_data):
    """Get contact details based on preferred method of contact"""
    preferred_method = csv_data.get('Preferred method of contact', '').lower()
    
    if 'home phone' in preferred_method:
        return csv_data.get('Home phone (Contact Details of the Client)', 'Home phone (Contact Details of the Client)')
    elif 'mobile' in preferred_method:
        return csv_data.get('Mobile phone (Contact Details of the Client)', 'Mobile phone (Contact Details of the Client)')
    elif 'work phone' in preferred_method:
        return csv_data.get('Work phone (Contact Details of the Client)', 'Work phone (Contact Details of the Client)')
    elif 'email' in preferred_method:
        return csv_data.get('Email address (Contact Details of the Client)', 'Email address (Contact Details of the Client)')
    else:
        # Default to home phone if no clear preference
        return csv_data.get('Home phone (Contact Details of the Client)', 'Home phone (Contact Details of the Client)')

# Signature extraction removed to prevent timeouts
def _extract_signatures_from_pdf_removed(source_pdf_path):
    """
    Extract signature images from the source PDF.
    Returns a dictionary with signature images (file paths).
    Signatures can be stored as:
    1. Form field signature values
    2. Embedded images in the PDF
    3. Annotations
    
    Uses multiple methods with fallbacks:
    - PyMuPDF (if available) - best for image extraction
    - pypdf - for form field signatures and embedded images
    - pdfplumber - for text-based identification
    
    This function is designed to fail gracefully - if extraction fails,
    it returns an empty dict and the PDF generation continues without signatures.
    
    OPTIMIZED: Only processes last 5 pages (where signatures usually are) to avoid timeouts.
    """
    import signal
    import time
    
    signatures = {}
    if not source_pdf_path or not os.path.exists(source_pdf_path):
        print(f"Signature extraction: Source PDF not found: {source_pdf_path}")
        return signatures
    
    print(f"Signature extraction: Attempting to extract from {source_pdf_path}")
    
    # Set a maximum processing time (30 seconds)
    start_time = time.time()
    MAX_PROCESSING_TIME = 30
    MAX_PAGES_TO_PROCESS = 5  # Only check last 5 pages (signatures are usually at the end)
    
    try:
        # Method 1: Try PyMuPDF first (best method, but optional)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(source_pdf_path)
            image_list = []
            
            total_pages = len(doc)
            print(f"Signature extraction: Using PyMuPDF, found {total_pages} pages")
            
            # Only process last MAX_PAGES_TO_PROCESS pages (signatures are usually at the end)
            start_page = max(0, total_pages - MAX_PAGES_TO_PROCESS)
            print(f"Signature extraction: Processing pages {start_page + 1} to {total_pages} (last {MAX_PAGES_TO_PROCESS} pages)")
            
            for page_num in range(start_page, total_pages):
                # Check timeout
                if time.time() - start_time > MAX_PROCESSING_TIME:
                    print("Signature extraction: Timeout reached during PyMuPDF processing")
                    break
                
                page = doc[page_num]
                # Get all images on the page
                page_images = page.get_images()
                print(f"Signature extraction: Page {page_num + 1} has {len(page_images)} images")
                # Limit images per page
                image_list.extend(page_images[:10])  # Max 10 images per page
            
            print(f"Signature extraction: Total images found: {len(image_list)}")
            
            # Extract images that might be signatures
            # Look for images at the bottom of pages (where signatures usually are)
            # Limit to first 20 images total to avoid memory issues
            for img_index, img in enumerate(image_list[:20]):
                # Check timeout
                if time.time() - start_time > MAX_PROCESSING_TIME:
                    print("Signature extraction: Timeout reached during image extraction")
                    break
                
                # Stop if we already have 2 signatures
                if len(signatures) >= 2:
                    break
                
                try:
                    # Get image data
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Save to temporary file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{image_ext}') as tmp_file:
                        tmp_file.write(image_bytes)
                        tmp_path = tmp_file.name
                    
                    print(f"Signature extraction: Extracted image {img_index + 1} ({len(image_bytes)} bytes) to {tmp_path}")
                    
                    # Take first 2 images as potential signatures
                    if img_index < 2:
                        key = 'signatory' if img_index == 0 else 'nc_representative'
                        signatures[key] = tmp_path
                        print(f"Signature extraction: Assigned image {img_index + 1} as {key}")
                except Exception as e:
                    print(f"Error extracting image {img_index}: {e}")
            
            doc.close()
            if signatures:
                print(f"Signature extraction: Successfully extracted {len(signatures)} signatures using PyMuPDF")
                return signatures
            else:
                print("Signature extraction: PyMuPDF found images but none were assigned as signatures")
        except ImportError:
            # PyMuPDF not available, continue to other methods
            print("Signature extraction: PyMuPDF not available, trying pypdf methods")
            pass
        except Exception as e:
            print(f"Error using PyMuPDF: {e}")
        
        # Method 2: Try to extract from signature form fields using pypdf
        if PdfReader is not None:
            try:
                reader = PdfReader(source_pdf_path)
                
                if reader.get_fields():
                    for field_name, field in reader.get_fields().items():
                        field_name_lower = (field_name or "").lower()
                        
                        # Check if this is a signature field
                        if 'signature' in field_name_lower or 'sign' in field_name_lower:
                            try:
                                # Try to extract signature image from appearance stream
                                if hasattr(field, 'get'):
                                    ap = field.get('/AP')
                                    if ap:
                                        normal_ap = ap.get('/N')
                                        if normal_ap:
                                            # Try to get image from appearance stream
                                            try:
                                                # Access the stream object
                                                if hasattr(normal_ap, 'get_data'):
                                                    stream_data = normal_ap.get_data()
                                                    # Try to extract image from PDF stream
                                                    # This is complex - would need PDF stream parsing
                                                    pass
                                                elif hasattr(normal_ap, 'get_object'):
                                                    # Try to get the object and extract image
                                                    obj = normal_ap.get_object()
                                                    if obj and hasattr(obj, 'get_data'):
                                                        stream_data = obj.get_data()
                                                        # Would need to parse PDF stream to extract image
                                                        pass
                                            except Exception as e:
                                                print(f"Error extracting from appearance stream: {e}")
                            except Exception as e:
                                print(f"Error extracting signature from field {field_name}: {e}")
            except Exception as e:
                print(f"Error reading PDF fields: {e}")
        
        # Method 3: Try using pdfplumber to extract images (better for FlateDecode images)
        if not signatures:
            try:
                # Check timeout
                if time.time() - start_time > MAX_PROCESSING_TIME:
                    print("Signature extraction: Timeout reached, skipping pdfplumber method")
                    return signatures
                
                print("Signature extraction: Trying pdfplumber image extraction method")
                import pdfplumber
                with pdfplumber.open(source_pdf_path) as pdf:
                    image_count = 0
                    total_images_found = 0
                    total_pages = len(pdf.pages)
                    
                    # Only process last MAX_PAGES_TO_PROCESS pages (signatures are usually at the end)
                    start_page = max(0, total_pages - MAX_PAGES_TO_PROCESS)
                    print(f"Signature extraction: Processing pages {start_page + 1} to {total_pages} (last {MAX_PAGES_TO_PROCESS} pages)")
                    
                    for page_num in range(start_page, total_pages):
                        # Check timeout before processing each page
                        if time.time() - start_time > MAX_PROCESSING_TIME:
                            print("Signature extraction: Timeout reached during page processing")
                            break
                        
                        page = pdf.pages[page_num]
                        images = page.images
                        if images:
                            print(f"Signature extraction: Page {page_num + 1} has {len(images)} images")
                            total_images_found += len(images)
                            
                            # Limit number of images processed per page
                            max_images_per_page = 10
                            images_to_process = images[:max_images_per_page]
                            
                            for img in images_to_process:
                                # Check timeout before processing each image
                                if time.time() - start_time > MAX_PROCESSING_TIME:
                                    print("Signature extraction: Timeout reached during image processing")
                                    break
                                
                                # Stop if we already have 2 signatures
                                if image_count >= 2:
                                    break
                                
                                try:
                                    # Check if this looks like a signature (usually at bottom of page, specific size)
                                    # Signatures are typically wider than tall, and positioned near bottom
                                    img_height = img.get('height', 0) or 0
                                    img_width = img.get('width', 0) or 0
                                    y_position = img.get('y0', 0) or 0
                                    page_height = getattr(page, 'height', None) or 800  # Default to 800 if None
                                    
                                    # Signatures are usually at bottom 20% of page and have reasonable dimensions
                                    is_likely_signature = (
                                        img_height and img_height > 20 and img_height < 200 and  # Reasonable signature height
                                        img_width and img_width > 100 and  # Signatures are usually wide
                                        page_height and y_position < page_height * 0.3  # Near bottom of page (y0 is from bottom)
                                    )
                                    
                                    if is_likely_signature or image_count < 2:  # Take first 2 images if we can't determine
                                        stream = img.get('stream')
                                        if stream:
                                            try:
                                                # Try to get raw image data
                                                if hasattr(stream, 'get_data'):
                                                    image_data = stream.get_data()
                                                elif hasattr(stream, '_data'):
                                                    image_data = stream._data
                                                else:
                                                    # Try to extract from the stream object
                                                    image_data = None
                                                    try:
                                                        # For FlateDecode, we need to decompress
                                                        import zlib
                                                        if hasattr(stream, 'raw_bytes'):
                                                            raw_bytes = stream.raw_bytes
                                                            # Try to decompress if it's compressed
                                                            try:
                                                                image_data = zlib.decompress(raw_bytes)
                                                            except:
                                                                image_data = raw_bytes
                                                    except Exception as e:
                                                        print(f"Error extracting stream data: {e}")
                                                
                                                if image_data and len(image_data) > 100:
                                                    # Determine file extension based on filter
                                                    suffix = '.png'  # Default for FlateDecode
                                                    if hasattr(stream, 'get'):
                                                        filter_type = stream.get('/Filter', '')
                                                        if '/DCTDecode' in (filter_type if isinstance(filter_type, list) else [filter_type]):
                                                            suffix = '.jpg'
                                                    
                                                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                                                        tmp_file.write(image_data)
                                                        tmp_path = tmp_file.name
                                                    
                                                    print(f"Signature extraction: Saved image from page {page_num + 1} ({len(image_data)} bytes) to {tmp_path}")
                                                    
                                                    if image_count < 2:
                                                        key = 'signatory' if image_count == 0 else 'nc_representative'
                                                        signatures[key] = tmp_path
                                                        image_count += 1
                                                        print(f"Signature extraction: Assigned image as {key}")
                                            except Exception as e:
                                                print(f"Error extracting image data: {e}")
                                                import traceback
                                                traceback.print_exc()
                                except Exception as e:
                                    print(f"Error processing image: {e}")
                    
                    print(f"Signature extraction: pdfplumber found {total_images_found} total images, extracted {len(signatures)} as signatures")
            except Exception as e:
                print(f"Error extracting images with pdfplumber: {e}")
                import traceback
                traceback.print_exc()
        
        # Method 4: Try using pypdf to extract images from pages (fallback)
        if not signatures and PdfReader is not None:
            try:
                print("Signature extraction: Trying pypdf image extraction method (fallback)")
                reader = PdfReader(source_pdf_path)
                image_count = 0
                total_images_found = 0
                
                total_pages = len(reader.pages)
                # Only process last MAX_PAGES_TO_PROCESS pages
                start_page = max(0, total_pages - MAX_PAGES_TO_PROCESS)
                print(f"Signature extraction: Processing pages {start_page + 1} to {total_pages} (last {MAX_PAGES_TO_PROCESS} pages)")
                
                for page_num in range(start_page, total_pages):
                    # Check timeout
                    if time.time() - start_time > MAX_PROCESSING_TIME:
                        print("Signature extraction: Timeout reached during pypdf processing")
                        break
                    
                    # Stop if we already have 2 signatures
                    if len(signatures) >= 2:
                        break
                    
                    try:
                        page = reader.pages[page_num]
                        resources = page.get('/Resources', {})
                        if resources and '/XObject' in resources:
                            xobjects = resources['/XObject']
                            if hasattr(xobjects, 'get_object'):
                                xobjects = xobjects.get_object()
                            
                            if isinstance(xobjects, dict):
                                for obj_name, obj in xobjects.items():
                                    try:
                                        if hasattr(obj, 'get') and obj.get('/Subtype') == '/Image':
                                            total_images_found += 1
                                            print(f"Signature extraction: Found image object {obj_name} on page {page_num + 1}")
                                            
                                            # Extract image data - try multiple filter types
                                            filter_type = obj.get('/Filter', '')
                                            is_supported = False
                                            
                                            # Check for JPEG (DCTDecode)
                                            if filter_type == '/DCTDecode' or (isinstance(filter_type, list) and '/DCTDecode' in filter_type):
                                                is_supported = True
                                                suffix = '.jpg'
                                            # Check for PNG (FlateDecode) - this is what signatures use!
                                            elif filter_type == '/FlateDecode' or (isinstance(filter_type, list) and '/FlateDecode' in filter_type):
                                                is_supported = True
                                                suffix = '.png'
                                            # Check for other common formats
                                            elif '/CCITTFaxDecode' in (filter_type if isinstance(filter_type, list) else [filter_type]):
                                                is_supported = True
                                                suffix = '.tiff'
                                            
                                            if is_supported:
                                                try:
                                                    image_data = obj.get_data()
                                                    if image_data and len(image_data) > 100:  # Only save if it's substantial
                                                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                                                            tmp_file.write(image_data)
                                                            tmp_path = tmp_file.name
                                                        
                                                        print(f"Signature extraction: Saved image {obj_name} ({len(image_data)} bytes) to {tmp_path}")
                                                        
                                                        if image_count < 2:
                                                            key = 'signatory' if image_count == 0 else 'nc_representative'
                                                            signatures[key] = tmp_path
                                                            image_count += 1
                                                            print(f"Signature extraction: Assigned image as {key}")
                                                except Exception as e:
                                                    print(f"Error writing image {obj_name} to temp file: {e}")
                                            else:
                                                print(f"Signature extraction: Image {obj_name} has unsupported filter: {filter_type}")
                                    except Exception as e:
                                        print(f"Error processing image object {obj_name}: {e}")
                    except Exception as e:
                        print(f"Error processing page {page_num}: {e}")
                        continue
                
                print(f"Signature extraction: pypdf found {total_images_found} total images, extracted {len(signatures)} as signatures")
            except Exception as e:
                print(f"Error extracting images with pypdf: {e}")
                import traceback
                traceback.print_exc()
                
    except Exception as e:
        print(f"Error extracting signatures: {e}")
        import traceback
        traceback.print_exc()
        # Return empty dict on error - don't break the build
    
    if signatures:
        print(f"Signature extraction: Successfully extracted {len(signatures)} signatures: {list(signatures.keys())}")
    else:
        print("Signature extraction: No signatures were extracted")
    
    return {}  # Always return empty - signature extraction disabled

def create_service_agreement_from_data(csv_data, output_path, contact_name=None, source_pdf_path=None, ndis_items=None, active_users=None):
    """
    Create a service agreement PDF from provided data dictionary.
    
    Args:
        csv_data: Dictionary containing form data
        output_path: Path where the PDF should be saved
        contact_name: Optional name to use for Key Contact lookup
        source_pdf_path: Optional path to source PDF for signature extraction
        ndis_items: Optional pre-loaded NDIS items (for performance)
        active_users: Optional pre-loaded active users (for performance)
    """
    # Load NDIS support items if not provided
    if ndis_items is None:
        ndis_items = load_ndis_support_items()
    
    # Get team value to determine which active users CSV to use
    team_value = csv_data.get('Neighbourhood Care representative team', '')
    # Clean up checkbox characters
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Load active users based on team if not provided
    if active_users is None:
        active_users = load_active_users(team_value)
    
    # Signature extraction removed to prevent timeouts
    signatures = {}
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    _build_service_agreement_content(doc, csv_data, ndis_items, active_users, contact_name, signatures)

def create_service_agreement():
    # Load NDIS support items
    ndis_items = load_ndis_support_items()
    
    # Read data from PDF (preferred) or CSV (fallback)
    csv_data = {}
    pdf_path = 'outputs/other/Neighbourhood Care Welcoming Form Template 2.pdf'
    
    # Try to parse PDF first
    if os.path.exists(pdf_path):
        try:
            csv_data = parse_pdf_to_data(pdf_path)
            print(f"Successfully parsed PDF: {pdf_path}")
        except Exception as e:
            print(f"Error parsing PDF: {e}. Falling back to CSV.")
            csv_data = {}
    
    # Fallback to CSV if PDF parsing failed or didn't get enough data
    if not csv_data or len([v for v in csv_data.values() if v]) < 5:
        csv_candidates = [
            'outputs/other/Neighbourhood Care Welcoming Form Template 2.csv',
            'outputs/other/Neighbourhood Care Welcoming Form.csv'
        ]
        last_err = None
        for candidate in csv_candidates:
            try:
                with open(candidate, 'r', encoding='utf-8') as file:
                    reader = csv.DictReader(file)
                    for row in reader:
                        csv_data = row
                        break
                print(f"Successfully loaded CSV: {candidate}")
                break
            except Exception as e:
                last_err = e
                continue
        if not csv_data and last_err:
            raise last_err
    
    # Get team value to determine which active users CSV to use
    team_value = csv_data.get('Neighbourhood Care representative team', '')
    # Clean up checkbox characters
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Load active users based on team
    active_users = load_active_users(team_value)
    
    # Create PDF document
    doc = SimpleDocTemplate("Service Agreement - FINAL TABLES.pdf", pagesize=A4)
    _build_service_agreement_content(doc, csv_data, ndis_items, active_users)

def _add_header_footer(canvas_obj, doc):
    """Add header and footer to PDF pages"""
    # Footer color (gray)
    footer_color = colors.HexColor('#7F7F7F')
    
    # Footer text
    footer_text = "Neighbourhood Care | Suite 103, 19 Ogilvie Road, Mount Pleasant, WA 6153 | ABN 40 634 832 607"
    
    # Get page number
    page_num = canvas_obj.getPageNumber()
    
    # Footer position - center text at normal footer position
    footer_y = 30  # Position for footer text
    
    # Draw footer text in center
    canvas_obj.saveState()
    canvas_obj.setFillColor(footer_color)
    canvas_obj.setFont("Helvetica", 8)
    
    # Center footer text
    page_width = A4[0]
    footer_x_center = page_width / 2
    canvas_obj.drawCentredString(footer_x_center, footer_y, footer_text)
    
    # Page number on right side (a bit below the footer text)
    page_num_text = str(page_num)
    page_num_y = footer_y - 12  # Position below the footer text
    page_num_x = page_width - 50  # Right side with margin
    canvas_obj.drawRightString(page_num_x, page_num_y, page_num_text)
    
    canvas_obj.restoreState()

def _add_first_page_header(canvas_obj, doc):
    """Add header with image to first page only"""
    # Add the image to the right side of header
    image_filename = 'image.png'
    
    # Get the script directory (where this file is located)
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
    except:
        script_dir = os.getcwd()
    
    # Search for the image file
    image_path = None
    search_dirs = [script_dir, os.getcwd(), '.']
    
    for search_dir in search_dirs:
        if os.path.exists(search_dir):
            try:
                # Try exact filename first
                test_path = os.path.join(search_dir, image_filename)
                if os.path.exists(test_path):
                    image_path = os.path.abspath(test_path)
                    print(f"DEBUG: Found image file: {image_filename}")
                    break
                
                # Fallback: search for any file with "image" in the name
                for filename in os.listdir(search_dir):
                    if filename.lower() == image_filename.lower() or (filename.lower().startswith('image') and filename.lower().endswith('.png')):
                        full_path = os.path.join(search_dir, filename)
                        if os.path.exists(full_path):
                            image_path = os.path.abspath(full_path)
                            print(f"DEBUG: Found image file: {filename}")
                            break
                if image_path:
                    break
            except Exception as e:
                print(f"DEBUG: Error searching in {search_dir}: {e}")
                continue
    
    print(f"DEBUG: Looking for image: {image_filename}")
    print(f"DEBUG: Image found: {image_path}")
    print(f"DEBUG: Image exists: {os.path.exists(image_path) if image_path else False}")
    
    if os.path.exists(image_path):
        try:
            # Image size - doubled from original size
            # Body text is 11pt, so image height around 80 points for better visibility
            img_height = 80  # Doubled size for better visibility
            
            # Position on right side of header
            # ReportLab uses bottom-left as origin (0,0), so y increases upward
            page_width = A4[0]  # 595.27 points
            page_height = A4[1]  # 841.89 points
            img_y = page_height - 70  # Top of page with margin (from bottom)
            
            # Try to get image dimensions to calculate aspect ratio
            img_width = img_height * 1.5  # Default aspect ratio
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.open(image_path)
                img_width_orig, img_height_orig = pil_img.size
                aspect_ratio = img_width_orig / img_height_orig
                img_width = img_height * aspect_ratio
                print(f"DEBUG: Image dimensions: {img_width_orig}x{img_height_orig}, scaled to: {img_width:.1f}x{img_height}")
            except Exception as pil_error:
                print(f"DEBUG: PIL not available, using default aspect ratio. Error: {pil_error}")
            
            img_x = page_width - img_width - 50  # Right side with margin
            print(f"DEBUG: Drawing image at position: x={img_x:.1f}, y={img_y}, size: {img_width:.1f}x{img_height}")
            print(f"DEBUG: Page dimensions: {page_width}x{page_height}")
            
            canvas_obj.saveState()
            # Use absolute path for drawImage
            abs_image_path = os.path.abspath(image_path)
            print(f"DEBUG: Using absolute image path: {abs_image_path}")
            
            # Draw the image
            canvas_obj.drawImage(
                abs_image_path, 
                img_x, 
                img_y, 
                width=img_width, 
                height=img_height, 
                preserveAspectRatio=True
            )
            canvas_obj.restoreState()
            print("DEBUG: Image drawn successfully")
        except Exception as e:
            print(f"ERROR: Could not add header image: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"ERROR: Image file not found at: {image_path}")
        print(f"DEBUG: Current working directory: {os.getcwd()}")
        print(f"DEBUG: Script directory: {os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else 'unknown'}")
    
    # Also add footer for first page
    _add_header_footer(canvas_obj, doc)

def _build_service_agreement_content(doc, csv_data, ndis_items, active_users, contact_name=None, signatures=None):
    """Build the service agreement PDF content"""
    story = []
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=18,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=14,
        leftIndent=0
    )
    
    # Style with no space after for immediate following text
    normal_no_space_style = ParagraphStyle(
        'CustomNormalNoSpace',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=14,
        leftIndent=0
    )
    
    # Style for headings that should have no space after
    heading_no_space_style = ParagraphStyle(
        'CustomHeadingNoSpace',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    
    # Style for black headings with no space after
    black_heading_no_space_style = ParagraphStyle(
        'BlackHeadingNoSpace',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.black,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    
    # Style for bold headings (questions) with no space after
    bold_heading_no_space_style = ParagraphStyle(
        'BoldHeadingNoSpace',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=14,
        leftIndent=0
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=12,
        leftIndent=0
    )
    
    black_heading_style = ParagraphStyle(
        'BlackHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.black,
        alignment=TA_LEFT,
        spaceAfter=12,
        leftIndent=0
    )
    
    table_text_style = ParagraphStyle(
        'TableText',
        wordWrap='CJK',  # Enable word wrapping
        parent=styles['Normal'],
        fontSize=8,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=10,
        leftIndent=0
    )
    
    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=20,
        bulletIndent=10,
        leading=14
    )
    
    # Get team value early for price determination
    team_value = csv_data.get('Neighbourhood Care representative team', '[To be filled in]')
    # Clean up checkbox characters that appear as black boxes
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Title
    story.append(Paragraph("Service Agreement", title_style))
    story.append(Spacer(1, 12))  # Add one line space after title
    
    # Introduction
    intro1 = "Thank you for choosing Neighbourhood Care. We look forward to working with you to help you achieve your goals."
    story.append(Paragraph(intro1, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    intro2 = "This document is a written agreement between you and Neighbourhood Care that outlines the supports we will provide and how they will be delivered."
    story.append(Paragraph(intro2, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    intro3 = "<b>Please make sure you have read and understood our Agreements, Promises and Terms of Service before completing this document.</b>"
    story.append(Paragraph(intro3, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    intro4 = "If you are unsure about any part of this document please speak to your Neighbourhood Care representative."
    story.append(Paragraph(intro4, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    intro5 = "This Service Agreement must then be signed for us to start delivering services."
    story.append(Paragraph(intro5, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    # What makes up your service
    what_makes_up_heading_style = ParagraphStyle(
        'WhatMakesUpHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("What makes up your service?", what_makes_up_heading_style))
    
    service_text = "Please note that your service is made up of face to face and some non face to face supports. Services that may be charged as part of your service are:"
    story.append(Paragraph(service_text, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    service_bullets = [
        "Transporting you during a shift (this is a $1 cost per km and is billed out of your core budget).",
        "Communication by phone or email or in a face to face meeting with key people in your network - when this is not part of your rostered shift.",
        "Travel for support workers or therapists when they are coming directly from the office or from another participant or travelling back to the office at the end of the shift.",
        "Preparing some reports that are required for the NDIS such as creating your Support Plan.",
        "Costs for when we are supporting you in the community such as parking, public transport and so forth.",
        "For <i>new</i> participants, receiving Core supports, the one off Establishment fee is applied."
    ]
    
    for bullet in service_bullets:
        story.append(Paragraph(f"• {bullet}", bullet_style))
    
    story.append(Spacer(1, 12))
    
    # Calculate Establishment Fee
    establishment_fee_amount = get_establishment_fee(csv_data, ndis_items, team_value)
    
    # Only show establishment fee table if fee is greater than $0.00
    if establishment_fee_amount and establishment_fee_amount != '$0.00':
        # Extract numeric value for comparison
        try:
            fee_value = float(establishment_fee_amount.replace('$', '').replace(',', ''))
            if fee_value > 0:
                # Establishment Fee
                establishment_data = [
                    ['Establishment Fee', establishment_fee_amount]
                ]
                
                establishment_table = Table(establishment_data, colWidths=[3*inch, 2*inch])
                establishment_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), BLUE_COLOR),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.white),
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('BACKGROUND', (1, 0), (1, 0), colors.white),
                    ('TEXTCOLOR', (1, 0), (1, 0), colors.black),
                    ('FONTNAME', (1, 0), (1, 0), 'Helvetica'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                ]))
                story.append(establishment_table)
                story.append(Spacer(1, 12))
        except (ValueError, AttributeError):
            # If we can't parse the fee, don't show the table
            pass
    
    # Schedule of Supports
    schedule_heading_style = ParagraphStyle(
        'ScheduleHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("Schedule of Supports", schedule_heading_style))
    
    # Core and Capacity Building
    core_capacity_heading_style = ParagraphStyle(
        'CoreCapacityHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.black,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("Core and Capacity Building", core_capacity_heading_style))
    # Create white bold text style for table cells
    white_bold_table_text_style = ParagraphStyle(
        'WhiteBoldTableText',
        parent=styles['Normal'],
        fontSize=8,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=10,
        leftIndent=0,
        textColor=colors.white,
        fontName='Helvetica-Bold'
    )
    
    core_data = [
        [Paragraph('Core Budget Allocated to Neighbourhood Care', white_bold_table_text_style), Paragraph(csv_data.get('Total core budget to allocate to Neighbourhood Care', 'Total core budget to allocate to Neighbourhood Care (NDIS Information)'), table_text_style)],
        [Paragraph('Capacity Building Budget Allocated to Neighbourhood Care', white_bold_table_text_style), Paragraph(csv_data.get('Total capacity building budget to allocate to Neighbourhood Care', 'Total capacity building budget to allocate to Neighbourhood Care (NDIS Information)'), table_text_style)]
    ]
    
    core_table = Table(core_data, colWidths=[3.5*inch, 2*inch])
    core_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), BLUE_COLOR),
        ('TEXTCOLOR', (0, 0), (0, 0), colors.white),
        ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
        ('BACKGROUND', (1, 0), (1, 0), colors.white),
        ('TEXTCOLOR', (1, 0), (1, 0), colors.black),
        ('FONTNAME', (1, 0), (1, 0), 'Helvetica'),
        ('BACKGROUND', (0, 1), (0, 1), BLUE_COLOR),
        ('TEXTCOLOR', (0, 1), (0, 1), colors.white),
        ('FONTNAME', (0, 1), (0, 1), 'Helvetica-Bold'),
        ('BACKGROUND', (1, 1), (1, 1), colors.white),
        ('TEXTCOLOR', (1, 1), (1, 1), colors.black),
        ('FONTNAME', (1, 1), (1, 1), 'Helvetica'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(core_table)
    
    # Support Items
    support_items_heading_style = ParagraphStyle(
        'SupportItemsHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.black,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("Support Items", support_items_heading_style))
    support_data = [['Category', 'Name', 'Number', 'Unit', 'Price']]
    
    # Extract support items from the PDF data - look for "Support item (X) (Support Items Required)"
    support_items_from_pdf = []
    for i in range(1, 20):  # Check up to 20 support items
        key = f'Support item ({i}) (Support Items Required)'
        item_name = csv_data.get(key, '').strip()
        if item_name:
            support_items_from_pdf.append((i, item_name))
    
    # If no support items found in PDF, use empty list (don't show hardcoded items)
    # Determine which state's price to use based on team
    price_state = get_price_state(team_value)
    price_key = 'wa_price' if price_state == 'WA' else 'qld_price'
    
    for item_num, item_name in support_items_from_pdf:
        item_details = lookup_support_item(ndis_items, item_name)
        # Check if item was actually found (not the placeholder)
        item_found = item_name in ndis_items or any(
            item_name.lower() in key.lower() or key.lower() in item_name.lower() 
            for key in ndis_items.keys()
        )
        # If item not found, show [Not Found] for all fields
        if item_found:
            support_data.append([
                Paragraph(f'Support item ({item_num})', table_text_style),
                Paragraph(item_name, table_text_style),
                Paragraph(item_details.get('number', ''), table_text_style),
                Paragraph(item_details.get('unit', ''), table_text_style),
                Paragraph(item_details.get(price_key, ''), table_text_style)
            ])
        else:
            support_data.append([
                Paragraph(f'Support item ({item_num})', table_text_style),
                Paragraph(item_name, table_text_style),
                Paragraph('[Not Found]', table_text_style),
                Paragraph('[Not Found]', table_text_style),
                Paragraph('[Not Found]', table_text_style)
            ])
    
    # Adjust column widths to prevent text overflow - A4 width is ~8.27 inches, leave some margin
    support_table = Table(support_data, colWidths=[0.7*inch, 3.5*inch, 1.1*inch, 0.7*inch, 0.9*inch])
    support_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), BLUE_COLOR),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(support_table)
    story.append(Spacer(1, 12))
    
    # Consents
    consents_heading_style = ParagraphStyle(
        'ConsentsHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("Consents", consents_heading_style))
    consent_data = []
    
    consents = [
        'I agree to receive services from Neighbourhood Care.',
        'I consent for Neighbourhood Care to create an NDIS portal service booking on my behalf if my budget/s are Agency Managed.',
        'I understand that if at any time I (The Participant) require emergency medical assistance, Neighbourhood Care staff will call an ambulance to attend, and that I (The Participant) will be liable for any expenses incurred for Ambulance attendance.',
        'I agree that Neighbourhood Care staff may administer simple first aid to me (The Participant), if the need arises.',
        'I consent for Neighbourhood Care to discuss relevant information about my case with other providers involved in my care and support, for example GP, support coordinator.',
        'I agree not to smoke inside the home whilst Neighbourhood Care staff are present.',
        'I understand that an Emergency Response Plan will be developed with me by Neighbourhood Care to help keep me safe in the event of an emergency.',
        'I consent for Neighbourhood Care for I (The Participant) to be photographed/recorded for therapeutic and/or training purposes.',
        'I give authority for my details or information to be shared with an external auditor who will assess Neighbourhood Care against the NDIS Quality and Safeguards Framework.'
    ]
    
    for consent in consents:
        consent_data.append([Paragraph(consent, white_bold_table_text_style), csv_data.get(consent, 'Yes')])
    
    consent_table = Table(consent_data, colWidths=[4.2*inch, 0.8*inch])
    consent_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), BLUE_COLOR),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),
        ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTSIZE', (0, 0), (-1, -1), 7),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(consent_table)
    story.append(Spacer(1, 12))
    
    # Agreements, Promises and Terms of Service
    agreements_heading_style = ParagraphStyle(
        'AgreementsHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("<b>Agreements, Promises and Terms of Service</b>", agreements_heading_style))
    
    story.append(Paragraph("Our Agreements, Promises and Terms of Service outline how we deliver services. It outlines our rights and responsibilities as a service provider, and the rights and responsibilities of the people we provide services to.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>What can you expect from Neighbourhood Care?</b>", bold_heading_no_space_style))
    story.append(Paragraph("We agree to:", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    nc_agreements = [
        "Review your care and service plan every 6 months with you.",
        "Maintain a service that works for you, so times of appointments meet your needs and we are in tune with each other. We call this Attunement.",
        "At all times communicate openly and honestly in a timely manner.",
        "At all times treat you with dignity and respect and being mindful of any cultural differences.",
        "Be open and transparent about managing complaints or disagreements and provide you the opportunity to provide feedback to us and to the NDIS.",
        "Ensure your privacy and any information is held in confidence and not shared without your permission.",
        "Work together at every step on your journey towards reaching your goals.",
        "Operate within the National Disability Insurance Scheme Act 2013 and associated Business Rules."
    ]
    
    for agreement in nc_agreements:
        story.append(Paragraph(f"• {agreement}", bullet_style))
    
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>What is expected of you as an NDIS participant?</b>", normal_style))
    story.append(Paragraph("You agree to:", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    participant_agreements = [
        "Inform Neighbourhood Care about how you wish your supports to be provided and how they should be offered to meet your needs.",
        "Treat Neighbourhood Care staff with courtesy and respect in the same way you want to be treated.",
        "Talk to Neighbourhood Care if you have any concerns about Plan Management or Financial Administration being provided.",
        "Give your care and support team the required notice if you need to end this Service Agreement. There is a notice period of 4 weeks to end this service.",
        "Advise your care and support team immediately if your plan is suspended or replaced by a new NDIS Plan or where you stop being an active participant in the NDIS."
    ]
    
    for agreement in participant_agreements:
        story.append(Paragraph(f"• {agreement}", bullet_style))
    
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>Cancellations</b>", normal_style))
    story.append(Paragraph("If a make-up shift with that support worker cannot be scheduled, the NDIS considers this a Short Notice Cancellation, and Neighbourhood Care may charge 100% of the agreed hourly rate.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>How will services be provided to you?</b>", bold_heading_no_space_style))
    story.append(Paragraph("Services will be provided at your place of residence and in other locations as deemed necessary and suitable by you, your family, the Support Coordinator and the Neighbourhood Care Team charged with your safety whilst in the service.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>When will services be provided?</b>", bold_heading_no_space_style))
    story.append(Paragraph("All services will be provided in attunement to your needs and subject to availability by others who may have an impact to your availability.", normal_no_space_style))
    story.append(Paragraph("From the commencement of the service agreement: Direct support provided as per the support and/or Therapy support plan, subject to change/increase, upon confirmation with you and/or your family.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>How long will services be provided?</b>", bold_heading_no_space_style))
    story.append(Paragraph("Services will be provided for the length of the service agreement plan unless otherwise ceased at the discretion by by you or your team, in accordance with Neighbourhood Care's Policy and Procedures.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>How to make changes?</b>", bold_heading_no_space_style))
    story.append(Paragraph("If changes to the supports or their delivery are required, you and your Neighbourhood Care team (the parties) agree to discuss and review this Service Agreement. The Parties agree that any changes to this Service Agreement will be in writing, signed, and dated by the Parties.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>How to end the Agreement?</b>", bold_heading_no_space_style))
    story.append(Paragraph("Should either Party wish to end this Service Agreement they must give 4 weeks written notice to their care and support team. If either Party seriously breaches this Service Agreement the requirement of notice will be waived.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>Pricing Changes</b>", bold_heading_no_space_style))
    story.append(Paragraph("Neighbourhood Care's services are charged in accordance with the NDIS Pricing Arrangements and Price Limits Guide. The prices set out in this Service Agreement will change in accordance with updates to the NDIS Pricing Arrangements and Price Limits Guide. This typically updates on the 1st of July each year but may be updated at other times.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>What to do if there is a problem?</b>", bold_heading_no_space_style))
    story.append(Paragraph("If there is a problem with anything related to your service or this agreement, you can contact:", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Your Neighbourhood Care contact person (please refer to the front page of your Service Agreement) or 1800 292 273.", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Alternatively, you can email your concern or query to: ask@nhcare.com.au.", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("If you don't feel that your problem was resolved please speak to your support coordinator, Local Area Coordinator or you can just contact the National Disability Insurance Agency (NDIA)", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>Collection of your personal information</b>", bold_heading_no_space_style))
    story.append(Paragraph("Neighbourhood Care will use your information to support your involvement in the NDIS.", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Neighbourhood Care will NOT use any of your personal information for any other purpose or disclose your personal information to any other organisations or individuals (including overseas recipients) unless authorised by law or you provide consent for us to do so.", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("You can also ask to see what personal information (if any) we hold about you at any time and you can seek correction if the information is incorrect.", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("<b>Neighbourhood Care's privacy policy describes:</b>", normal_no_space_style))
    
    privacy_bullets = [
        "How we use your personal information",
        "Why some personal information may be given to other organisations from time to time",
        "How you can access the personal information we have about you on our system",
        "How you can complain about a privacy breach, and how Neighbourhood Care deals with the complaint.",
        "How you can get your personal information corrected if it is wrong."
    ]
    
    for bullet in privacy_bullets:
        story.append(Paragraph(f"• {bullet}", bullet_style))
    
    story.append(Spacer(1, 12))
    story.append(Paragraph("You can find the policy by enquiring at Neighbourhood Care.", normal_no_space_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Please note that Neighbourhood Care is required to release information about service users (without identifying you by full name or address) to the Australian Institute of Health and Welfare, to enable statistics about disability services and their clients to be compiled. This information will be kept confidential. This information is used for statistical purposes only and will not be used to affect your entitlements or your access to services. You have the right to access your own files and to update or correct information included in the Disability Services National Minimum Data Set collection.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("<b>Goods and Services Tax</b>", normal_style))
    story.append(Paragraph("Most services provided under the NDIS will not include GST. However, GST will apply to some services. Neighbourhood Care will apply GST when it is required.", normal_no_space_style))
    story.append(Spacer(1, 12))
    
    # For your information text (comes before Signatures heading)
    for_your_info_text = 'For your information: "A supply of supports under this Service Agreement is a supply of one or more reasonable and necessary supports specified in the statement of supports included, under subsection 33(2) of the National Disability Insurance Scheme Act 2013 (NDIS Act), in the participant\'s NDIS Plan currently in effect under section 37 of the NDIS Act."'
    story.append(Paragraph(for_your_info_text, normal_no_space_style))
    story.append(Spacer(1, 12))
    
    # Signatures
    signatures_heading_style = ParagraphStyle(
        'SignaturesHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0
    )
    story.append(Paragraph("Signatures", signatures_heading_style))
    signatory_name = f"{csv_data.get('First name (Person Signing the Agreement)', 'First name (Person Signing the Agreement)')} {csv_data.get('Surname (Person Signing the Agreement)', 'Surname (Person Signing the Agreement)')}"
    signatory_text = f"<b>Signatory:</b><br/><b>Name:</b> {signatory_name}<br/><b>Date:</b> [Date]<br/><b>Signed:</b>"
    story.append(Paragraph(signatory_text, normal_no_space_style))
    
    # Add signatory signature image if available
    if signatures and 'signatory' in signatures:
        try:
            sig_img = signatures['signatory']
            # Check if it's a file path (string) or already an image object
            if isinstance(sig_img, str) and os.path.exists(sig_img):
                print(f"Adding signatory signature from: {sig_img}")
                story.append(Spacer(1, 6))
                story.append(Image(sig_img, width=2*inch, height=0.5*inch))
            else:
                print(f"Signatory signature path invalid or doesn't exist: {sig_img}")
                story.append(Paragraph("[Signature]", normal_no_space_style))
        except Exception as e:
            print(f"Error adding signatory signature: {e}")
            import traceback
            traceback.print_exc()
            story.append(Paragraph("[Signature]", normal_no_space_style))
    else:
        print(f"Signatory signature not found. Available signatures: {list(signatures.keys()) if signatures else 'none'}")
        story.append(Paragraph("[Signature]", normal_no_space_style))
    
    story.append(Spacer(1, 12))
    
    # Neighbourhood Care Representative
    nc_rep_text = f"<b>Neighbourhood Care Representative:</b><br/><b>Name:</b> [To be filled with NC representative name]<br/><b>Date:</b> [Date]<br/><b>Signed:</b>"
    story.append(Paragraph(nc_rep_text, normal_no_space_style))
    
    # Add NC representative signature image if available
    if signatures and 'nc_representative' in signatures:
        try:
            sig_img = signatures['nc_representative']
            # Check if it's a file path (string) or already an image object
            if isinstance(sig_img, str) and os.path.exists(sig_img):
                print(f"Adding NC representative signature from: {sig_img}")
                story.append(Spacer(1, 6))
                story.append(Image(sig_img, width=2*inch, height=0.5*inch))
            else:
                print(f"NC representative signature path invalid or doesn't exist: {sig_img}")
                story.append(Paragraph("[Signature]", normal_no_space_style))
        except Exception as e:
            print(f"Error adding NC representative signature: {e}")
            import traceback
            traceback.print_exc()
            story.append(Paragraph("[Signature]", normal_no_space_style))
    else:
        print(f"NC representative signature not found. Available signatures: {list(signatures.keys()) if signatures else 'none'}")
        story.append(Paragraph("[Signature]", normal_no_space_style))
    
    story.append(Spacer(1, 12))
    
    # Participant - FIXED with all missing fields
    story.append(Paragraph("Appendix", black_heading_style))
    story.append(Paragraph("Participant", black_heading_no_space_style))
    # Participant Name: First name + Middle name + Surname (from Details of the Client)
    first_name = csv_data.get('First name (Details of the Client)', '').strip()
    middle_name = csv_data.get('Middle name (Details of the Client)', '').strip()
    surname = csv_data.get('Surname (Details of the Client)', '').strip()
    participant_name_parts = [p for p in [first_name, middle_name, surname] if p]
    participant_name = ' '.join(participant_name_parts) if participant_name_parts else ''
    
    # Emergency Contact: First name + Surname (from Emergency contact)
    emergency_first = csv_data.get('First name (Emergency contact)', '').strip()
    emergency_surname = csv_data.get('Surname (Emergency contact)', '').strip()
    emergency_contact_parts = [p for p in [emergency_first, emergency_surname] if p]
    emergency_contact = ' '.join(emergency_contact_parts) if emergency_contact_parts else get_emergency_contact(csv_data)
    
    # Get all values, using empty string if not found
    home_address = csv_data.get('Home address (Contact Details of the Client)', '').strip()
    home_phone = csv_data.get('Home phone (Contact Details of the Client)', '').strip()
    mobile_phone = csv_data.get('Mobile phone (Contact Details of the Client)', '').strip()
    email_address = csv_data.get('Email address (Contact Details of the Client)', '').strip()
    dob = csv_data.get('Date of birth (Details of the Client)', '').strip() or csv_data.get('Date of birth', '').strip()
    ndis_num = csv_data.get('NDIS number (Details of the Client)', '').strip() or csv_data.get('NDIS number', '').strip()
    plan_start = csv_data.get('Plan start date', '').strip()
    plan_end = csv_data.get('Plan end date', '').strip()
    service_start = csv_data.get('Service start date', '').strip() or csv_data.get('Service start', '').strip()
    service_end = csv_data.get('Service end date', '').strip() or csv_data.get('Service end', '').strip()
    preferred_contact = csv_data.get('Preferred method of contact', '').strip()
    # Clean up checkbox characters that appear as black boxes
    preferred_contact = preferred_contact.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    participant_data = [
        ['Participant Name', Paragraph(participant_name, table_text_style)],
        ['Date of Birth', dob],
        ['NDIS Number', ndis_num],
        ['Plan Duration', f"{plan_start} - {plan_end}" if (plan_start or plan_end) else ''],
        ['Address', Paragraph(home_address, table_text_style)],
        ['Home phone', home_phone],
        ['Mobile phone', mobile_phone],
        ['Email address', Paragraph(email_address, table_text_style)],
        ['Preferred contact method', preferred_contact],
        ['Emergency Contact', Paragraph(emergency_contact, table_text_style)],
        ['Service Agreement Duration', f"{service_start} - {service_end}" if (service_start or service_end) else '']
    ]
    
    participant_table = Table(participant_data, colWidths=[2.5*inch, 3*inch])
    participant_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(participant_table)
    story.append(Spacer(1, 12))
    
    # Signatory (detailed) - FIXED with all missing fields
    story.append(Paragraph("Signatory", black_heading_no_space_style))
    # Get signatory contact details based on who is signing (preferred method only)
    signatory_contact = get_signatory_contact_details(csv_data)
    
    signatory_detailed_data = [
        ['Name', Paragraph(get_signatory_name(csv_data), table_text_style)],
        ['Relationship to Participant', get_signatory_relationship(csv_data)],
        ['Address', Paragraph(get_signatory_address(csv_data), table_text_style)],
        ['Contact Details', Paragraph(signatory_contact, table_text_style)]
    ]
    
    signatory_detailed_table = Table(signatory_detailed_data, colWidths=[2.5*inch, 3*inch])
    signatory_detailed_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(signatory_detailed_table)
    story.append(Spacer(1, 12))
    
    # Plan Manager
    story.append(Paragraph("Plan Manager", black_heading_no_space_style))
    plan_manager_data = [
        ['Name', get_plan_manager_name(csv_data)],
        ['Postal Address', Paragraph(get_plan_manager_address(csv_data), table_text_style)],
        ['Phone', get_plan_manager_phone(csv_data)],
        ['Email Address', Paragraph(get_plan_manager_email(csv_data), table_text_style)]
    ]
    
    plan_manager_table = Table(plan_manager_data, colWidths=[2.5*inch, 3*inch])
    plan_manager_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(plan_manager_table)
    story.append(Spacer(1, 12))
    
    # My Neighbourhood Care Key Contact
    story.append(Paragraph("My Neighbourhood Care Key Contact", black_heading_no_space_style))
    
    # Get contact name from parameter or fallback to Respondent field
    contact_name_to_use = contact_name or csv_data.get('Respondent', '')
    user_data = lookup_user_data(active_users, contact_name_to_use) if contact_name_to_use else {'name': '', 'mobile': '', 'email': ''}
    
    # Calculate My Neighbourhood Care ID: First name + Surname + Year of Date of birth
    first_name = csv_data.get('First name (Details of the Client)', '').strip()
    surname = csv_data.get('Surname (Details of the Client)', '').strip()
    dob = csv_data.get('Date of birth (Details of the Client)', '').strip()
    
    # Extract year from date of birth (handle formats like DD/MM/YYYY or YYYY-MM-DD)
    year = ''
    if dob:
        # Try to extract year from common date formats
        year_match = re.search(r'\b(19|20)\d{2}\b', dob)
        if year_match:
            year = year_match.group(0)
    
    # Build ID: First name + Surname + Year (with spaces)
    name_parts = [p for p in [first_name, surname] if p]
    neighbourhood_care_id = ' '.join(name_parts) + ' ' + year if name_parts and year else '[To be filled in]'
    
    # Get team value and clean checkbox characters
    team_value = csv_data.get('Neighbourhood Care representative team', '[To be filled in]')
    # Clean up checkbox characters that appear as black boxes
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    key_contact_data = [
        ['My Neighbourhood Care ID', neighbourhood_care_id],
        ['Team', team_value],
        ['Key Contact', Paragraph(contact_name_to_use if contact_name_to_use else user_data.get('name', '[To be filled in]'), table_text_style)],
        ['Phone', user_data.get('mobile', '[To be filled in]')],
        ['Email Address', Paragraph(user_data.get('email', '[To be filled in]'), table_text_style)],
        ['Neighbourhood Care Office', 'Phone: 1800 292 273']
    ]
    
    key_contact_table = Table(key_contact_data, colWidths=[2.5*inch, 3*inch])
    key_contact_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    story.append(key_contact_table)
    
    # Build PDF with headers and footers
    doc.build(story, onFirstPage=_add_first_page_header, onLaterPages=_add_header_footer)
    print("Service Agreement PDF FINAL TABLES created successfully!")

def get_emergency_contact(csv_data):
    """Get emergency contact based on the logic specified"""
    is_primary_carer = csv_data.get('Is the primary carer also the emergency contact for the participant?', '').strip()
    # Clean checkbox characters and check if it's "yes"
    is_primary_carer_clean = is_primary_carer.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip().lower()
    
    if 'yes' in is_primary_carer_clean:
        first_name = csv_data.get('First name (Primary carer)', '').strip()
        surname = csv_data.get('Surname (Primary carer)', '').strip()
        name_parts = [p for p in [first_name, surname] if p]
        return ' '.join(name_parts) if name_parts else ''
    else:
        first_name = csv_data.get('First name (Emergency contact)', '').strip()
        surname = csv_data.get('Surname (Emergency contact)', '').strip()
        name_parts = [p for p in [first_name, surname] if p]
        return ' '.join(name_parts) if name_parts else ''

def get_signatory_name(csv_data):
    """Get signatory name based on who is signing"""
    person_signing = csv_data.get('Person signing the agreement', '').strip()
    if person_signing.lower() == 'participant':
        # Participant is the client - use First name + Middle name + Surname from Details of the Client
        first_name = csv_data.get('First name (Details of the Client)', '').strip()
        middle_name = csv_data.get('Middle name (Details of the Client)', '').strip()
        surname = csv_data.get('Surname (Details of the Client)', '').strip()
        name_parts = [p for p in [first_name, middle_name, surname] if p]
        return ' '.join(name_parts) if name_parts else ''
    elif person_signing.lower() == 'primary carer':
        first_name = csv_data.get('First name (Primary carer)', '').strip()
        surname = csv_data.get('Surname (Primary carer)', '').strip()
        name_parts = [p for p in [first_name, surname] if p]
        return ' '.join(name_parts) if name_parts else ''
    else:
        first_name = csv_data.get('First name (Person Signing the Agreement)', '').strip()
        surname = csv_data.get('Surname (Person Signing the Agreement)', '').strip()
        name_parts = [p for p in [first_name, surname] if p]
        return ' '.join(name_parts) if name_parts else ''

def get_signatory_relationship(csv_data):
    """Get signatory relationship based on who is signing"""
    person_signing = csv_data.get('Person signing the agreement', '').strip()
    if person_signing.lower() == 'participant':
        return 'Participant'
    elif person_signing.lower() == 'primary carer':
        return csv_data.get('Relationship to client (Primary carer)', '').strip()
    else:
        return csv_data.get('Relationship to client (Person Signing the Agreement)', '').strip()

def get_signatory_address(csv_data):
    """Get signatory address based on who is signing"""
    person_signing = csv_data.get('Person signing the agreement', '').strip()
    if person_signing.lower() == 'participant':
        return csv_data.get('Home address (Contact Details of the Client)', '').strip()
    elif person_signing.lower() == 'primary carer':
        return csv_data.get('Home address (Primary carer)', '').strip()
    else:
        return csv_data.get('Home address (Person Signing the Agreement)', '').strip()

def get_signatory_contact_details(csv_data):
    """Get actual contact detail value for signatory based on preferred method and who is signing"""
    person_signing = csv_data.get('Person signing the agreement', '').strip()
    
    # Get preferred method of contact
    preferred_contact = ''
    if person_signing.lower() == 'participant':
        # Use participant's preferred method of contact
        preferred_contact = csv_data.get('Preferred method of contact', '').strip()
    elif person_signing.lower() == 'primary carer':
        # Use primary carer's preferred method of contact (if available)
        preferred_contact = csv_data.get('Preferred method of contact (Primary carer)', '').strip() or csv_data.get('Preferred method of contact', '').strip()
    else:
        # Use Person Signing the Agreement preferred method of contact (if available)
        preferred_contact = csv_data.get('Preferred method of contact (Person Signing the Agreement)', '').strip() or csv_data.get('Preferred method of contact', '').strip()
    
    # Clean up checkbox characters to get the actual preferred method
    if preferred_contact:
        preferred_contact = preferred_contact.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Get the actual contact value based on preferred method
    preferred_contact_lower = preferred_contact.lower()
    
    if person_signing.lower() == 'participant':
        # Get participant's contact details
        if 'home phone' in preferred_contact_lower:
            return csv_data.get('Home phone (Contact Details of the Client)', '').strip()
        elif 'mobile phone' in preferred_contact_lower or 'mobile' in preferred_contact_lower:
            return csv_data.get('Mobile phone (Contact Details of the Client)', '').strip()
        elif 'email' in preferred_contact_lower:
            return csv_data.get('Email address (Contact Details of the Client)', '').strip()
        elif 'work phone' in preferred_contact_lower:
            return csv_data.get('Work phone (Contact Details of the Client)', '').strip()
    elif person_signing.lower() == 'primary carer':
        # Get primary carer's contact details
        if 'home phone' in preferred_contact_lower:
            return csv_data.get('Home phone (Primary carer)', '').strip()
        elif 'mobile phone' in preferred_contact_lower or 'mobile' in preferred_contact_lower:
            return csv_data.get('Mobile phone (Primary carer)', '').strip()
        elif 'email' in preferred_contact_lower:
            return csv_data.get('Email address (Primary carer)', '').strip()
        elif 'work phone' in preferred_contact_lower:
            return csv_data.get('Work phone (Primary carer)', '').strip()
    else:
        # Get Person Signing the Agreement's contact details
        if 'home phone' in preferred_contact_lower:
            return csv_data.get('Home phone (Person Signing the Agreement)', '').strip()
        elif 'mobile phone' in preferred_contact_lower or 'mobile' in preferred_contact_lower:
            return csv_data.get('Mobile phone (Person Signing the Agreement)', '').strip()
        elif 'email' in preferred_contact_lower:
            return csv_data.get('Email address (Person Signing the Agreement)', '').strip()
        elif 'work phone' in preferred_contact_lower:
            return csv_data.get('Work phone (Person Signing the Agreement)', '').strip()
    
    # Fallback: return preferred method if we can't find the actual value
    return preferred_contact

def get_plan_manager_name(csv_data):
    """Get plan manager name based on plan management type"""
    plan_type = csv_data.get('Plan management type', '')
    if plan_type in ['NDIA Agency Managed', 'Insurance Commission of WA']:
        return ''
    else:
        return csv_data.get('Plan manager name', 'Plan manager name (Support Items Required)')

def get_plan_manager_address(csv_data):
    """Get plan manager address based on plan management type"""
    plan_type = csv_data.get('Plan management type', '')
    if plan_type in ['NDIA Agency Managed', 'Insurance Commission of WA']:
        return ''
    else:
        return csv_data.get('Plan manager postal address', 'Plan manager postal address (Support Items Required)')

def get_plan_manager_phone(csv_data):
    """Get plan manager phone based on plan management type"""
    plan_type = csv_data.get('Plan management type', '')
    if plan_type in ['NDIA Agency Managed', 'Insurance Commission of WA']:
        return ''
    else:
        return csv_data.get('Plan manager phone number', 'Plan manager phone number (Support Items Required)')

def get_plan_manager_email(csv_data):
    """Get plan manager email based on plan management type"""
    plan_type = csv_data.get('Plan management type', '')
    if plan_type in ['NDIA Agency Managed', 'Insurance Commission of WA']:
        return ''
    else:
        return csv_data.get('Plan manager email address', 'Plan manager email address (Support Items Required)')

def format_date_for_display(date_str):
    """Format date string to DD/MM/YYYY format"""
    if not date_str:
        return ""
    
    date_str = date_str.strip()
    if not date_str:
        return ""
    
    from datetime import datetime
    
    # Try to parse common date formats
    date_formats = [
        '%Y-%m-%d',      # 2023-12-25
        '%d/%m/%Y',      # 25/12/2023
        '%m/%d/%Y',      # 12/25/2023
        '%d-%m-%Y',      # 25-12-2023
        '%Y/%m/%d',      # 2023/12/25
        '%d.%m.%Y',      # 25.12.2023
    ]
    
    for fmt in date_formats:
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime('%d/%m/%Y')
        except ValueError:
            continue
    
    # If no format matched, try to extract numbers
    try:
        numbers = re.findall(r'\d+', date_str)
        if len(numbers) >= 3:
            if len(numbers[0]) == 4:
                year, month, day = numbers[0], numbers[1], numbers[2]
            elif len(numbers[-1]) == 4:
                day, month, year = numbers[0], numbers[1], numbers[2]
            else:
                day, month, year = numbers[0], numbers[1], numbers[2]
            
            day = day.zfill(2)
            month = month.zfill(2)
            return f"{day}/{month}/{year}"
    except (ValueError, IndexError, AttributeError):
        pass
    
    return date_str

def get_emergency_contact_phone(csv_data):
    """Get emergency contact phone numbers (Home phone + Mobile phone + Work phone)"""
    phones = []
    
    # ONLY get emergency contact phone fields - no fallback to primary carer
    home_phone = csv_data.get('Home phone (Emergency contact)', '').strip()
    mobile_phone = csv_data.get('Mobile phone (Emergency contact)', '').strip()
    work_phone = csv_data.get('Work phone (Emergency contact)', '').strip()
    
    # Clean phone numbers of any special characters that might cause rendering issues
    def clean_phone(phone):
        if not phone:
            return ''
        # Remove any special unicode characters that might render as black squares
        cleaned = phone.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '')
        # Remove semicolons and any other problematic characters
        cleaned = cleaned.replace(';', '').replace(',', '')  # Remove semicolons and commas from phone numbers
        # Keep only printable characters and common phone characters
        cleaned = ''.join(c for c in cleaned if c.isprintable() or c in [' ', '-', '(', ')', '+'])
        return cleaned.strip()
    
    if home_phone:
        phones.append(clean_phone(home_phone))
    if mobile_phone:
        phones.append(clean_phone(mobile_phone))
    if work_phone:
        phones.append(clean_phone(work_phone))
    
    # Join with semicolons (but semicolons are removed from individual phone numbers)
    return '; '.join(phones) if phones else ''

def get_emergency_contact_relationship(csv_data):
    """Get emergency contact relationship to client"""
    # ONLY get from emergency contact field - no fallback to primary carer
    relationship = csv_data.get('Relationship to client (Emergency contact)', '').strip()
    
    # If not found, return empty - no fallback searches
    return relationship if relationship else ''

def get_client_phone_numbers(csv_data):
    """Get client phone numbers (Home phone + Mobile phone + Work phone)"""
    phones = []
    
    home_phone = csv_data.get('Home phone (Contact Details of the Client)', '').strip()
    mobile_phone = csv_data.get('Mobile phone (Contact Details of the Client)', '').strip()
    work_phone = csv_data.get('Work phone (Contact Details of the Client)', '').strip()
    
    if home_phone:
        phones.append(home_phone)
    if mobile_phone:
        phones.append(mobile_phone)
    if work_phone:
        phones.append(work_phone)
    
    return '; '.join(phones) if phones else ''

def create_emergency_disaster_plan_from_data(csv_data, output_path, contact_name=None, active_users=None):
    """
    Create an Emergency & Disaster Plan PDF from provided data dictionary.
    
    Args:
        csv_data: Dictionary containing form data
        output_path: Path where the PDF should be saved
        contact_name: Optional name to use for Team member lookup
        active_users: Optional pre-loaded active users (for performance)
    """
    # Get team value to determine which active users CSV to use
    team_value = csv_data.get('Neighbourhood Care representative team', '')
    # Clean up checkbox characters
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Load active users based on team if not provided
    if active_users is None:
        active_users = load_active_users(team_value)
    
    # Get team member name (similar to Key Contact lookup)
    team_member_name_to_use = contact_name or csv_data.get('Respondent', '')
    team_member_data = lookup_user_data(active_users, team_member_name_to_use) if team_member_name_to_use else {'name': '', 'mobile': '', 'email': ''}
    # Use the name from lookup if available, otherwise use the provided name
    team_member_name = team_member_name_to_use if team_member_name_to_use else team_member_data.get('name', '')
    # If we have lookup data, prefer the name from the lookup (similar to Key Contact)
    if team_member_data.get('name') and team_member_data.get('name') != '[Not Found]':
        team_member_name = team_member_data.get('name', team_member_name)
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=12,
        leftIndent=0
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=8,
        spaceBefore=12,
        leftIndent=0
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=6,
        leading=14,
        leftIndent=0
    )
    
    table_text_style = ParagraphStyle(
        'TableText',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=12,
        leftIndent=0
    )
    
    # Title
    story.append(Paragraph("Emergency and Disaster Plan for Participants", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Introductory text
    intro_text = ("This Emergency and Disaster plan is to assist you to understand potential risks and how to protect "
                  "yourself and the people who support you. Your Neighbourhood Care Support Team will assist you to fill "
                  "this form out. Please refer to other relevant plans before completing this plan; such as support plan, "
                  "risk assessment, individual COVID-19 response plan & mediation assistance plan (if applicable).")
    story.append(Paragraph(intro_text, normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Section 1 heading
    story.append(Paragraph("1. Contact Information", heading_style))
    
    # General Information table
    first_name = csv_data.get('First name (Details of the Client)', '').strip()
    surname = csv_data.get('Surname (Details of the Client)', '').strip()
    client_name = ' '.join([p for p in [first_name, surname] if p]).strip() or ''
    client_phone = get_client_phone_numbers(csv_data)
    
    general_info_data = [
        [Paragraph("<b>General Information</b>", ParagraphStyle('TableHeading', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER))],
        ['Your name', Paragraph(client_name, table_text_style)],
        ['Your phone number', Paragraph(client_phone, table_text_style)]
    ]
    
    general_info_table = Table(general_info_data, colWidths=[2.5*inch, 3.5*inch])
    general_info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), BLUE_COLOR),  # Heading row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('SPAN', (0, 0), (-1, 0))  # Span heading across all columns
    ]))
    
    story.append(general_info_table)
    story.append(Spacer(1, 0.2*inch))
    
    # Key Emergency Contacts table
    # ONLY use fields from Emergency contact section - no fallback to primary carer
    emergency_first = csv_data.get('First name (Emergency contact)', '').strip()
    emergency_surname = csv_data.get('Surname (Emergency contact)', '').strip()
    
    emergency_name = ' '.join([p for p in [emergency_first, emergency_surname] if p]).strip() or ''
    emergency_phone = get_emergency_contact_phone(csv_data)
    emergency_relationship = get_emergency_contact_relationship(csv_data)
    
    # Ensure phone is a clean plain string (not Paragraph) - strip all special characters
    emergency_phone_clean = str(emergency_phone) if emergency_phone else ''
    # Remove any problematic unicode characters that render as black squares
    emergency_phone_clean = emergency_phone_clean.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '')
    # Remove semicolons - they shouldn't be in phone numbers
    emergency_phone_clean = emergency_phone_clean.replace(';', '')
    # Keep only printable ASCII characters and common phone characters (no semicolons)
    emergency_phone_clean = ''.join(c for c in emergency_phone_clean if (c.isprintable() and ord(c) < 128) or c in [' ', '-', '(', ')', '+', ','])
    emergency_phone_clean = emergency_phone_clean.strip()
    
    # Ensure relationship is displayed correctly
    emergency_relationship_clean = emergency_relationship if emergency_relationship else ''
    
    emergency_contacts_data = [
        [Paragraph("<b>Key Emergency Contacts</b>", ParagraphStyle('TableHeading', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER))],
        [Paragraph("<b>Name</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER)),
         Paragraph("<b>Phone</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER)),
         Paragraph("<b>Relationship</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER))],
        [Paragraph(emergency_name, table_text_style) if emergency_name else '', 
         emergency_phone_clean,  # Plain string, not Paragraph - already cleaned
         Paragraph(emergency_relationship_clean, table_text_style) if emergency_relationship_clean else '']
    ]
    
    emergency_contacts_table = Table(emergency_contacts_data, colWidths=[2*inch, 2*inch, 2*inch])
    emergency_contacts_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), BLUE_COLOR),  # Heading row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, 1), BLUE_COLOR),  # Header row
        ('TEXTCOLOR', (0, 1), (-1, 1), colors.white),
        ('BACKGROUND', (0, 2), (-1, 2), colors.white),  # Data row
        ('TEXTCOLOR', (0, 2), (-1, 2), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (-1, 1), 'CENTER'),
        ('ALIGN', (0, 2), (-1, 2), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('SPAN', (0, 0), (-1, 0))  # Span heading across all columns
    ]))
    
    story.append(emergency_contacts_table)
    story.append(Spacer(1, 0.2*inch))
    
    # My Important Contacts table
    important_contacts_data = [
        [Paragraph("<b>My Important Contacts</b>", ParagraphStyle('TableHeading', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER))],
        ['Advocate', ''],
        ['Power of Attorney/Guardian', ''],
        ['Solicitor', ''],
        ['Insurer (home)', ''],
        ['Insurer (vehicle)', ''],
        ['Childcare/School Contact', ''],
        ['Workplace/Volunteer Contact', ''],
        ['Doctor', ''],
        ['Specialist Practitioner', ''],
        ['Private Health Cover', '']
    ]
    
    important_contacts_table = Table(important_contacts_data, colWidths=[2.5*inch, 3.5*inch])
    important_contacts_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), BLUE_COLOR),  # Heading row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, 0), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('SPAN', (0, 0), (-1, 0))  # Span heading across all columns
    ]))
    
    story.append(important_contacts_table)
    story.append(Spacer(1, 0.2*inch))
    
    # What are the main risks in your community table
    risks = [
        'Heatwave', 'Storm', 'Cyclone', 'Bushfire', 'Flood', 'Earthquake',
        'Landslide', 'Tsunami', 'Assault', 'Power outage', 'Gas outage',
        'Health emergency', 'House fire', 'Burglary/break-in'
    ]
    
    risks_data = []
    for risk in risks:
        # Use plain ASCII characters for checkbox - avoid any unicode that might render as black squares
        risks_data.append([risk, '[ ]'])  # Simple ASCII brackets - no unicode characters
    
    risks_table = Table(risks_data, colWidths=[2.5*inch, 3.5*inch])
    risks_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(Paragraph("2. Identify Risks", heading_style))
    story.append(Paragraph("<i>What are the main risks in your community?</i>", normal_style))
    story.append(risks_table)
    story.append(Spacer(1, 0.2*inch))
    
    # How would the emergency affect you? table
    emergency_types = [
        'Heatwave', 'Storm', 'Cyclone', 'Bushfire', 'Flood', 'Earthquake',
        'Landslide', 'Tsunami', 'Assault', 'Power outage', 'Gas outage',
        'Health emergency', 'House fire', 'Burglary/break-in'
    ]
    
    emergency_affect_data = [
        [Paragraph("<b>Emergency Type</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER)),
         Paragraph("<b>How you're affected</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white, alignment=TA_CENTER))]
    ]
    
    # Add rows for each emergency type - make them bold
    for emergency_type in emergency_types:
        emergency_affect_data.append([Paragraph(f"<b>{emergency_type}</b>", table_text_style), ''])
    
    # Define gray color for backgrounds
    GRAY_COLOR = colors.HexColor('#d9d9d9')
    
    emergency_affect_table = Table(emergency_affect_data, colWidths=[2.5*inch, 3.5*inch])
    emergency_affect_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), BLUE_COLOR),  # Header row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (0, -1), GRAY_COLOR),  # Left column (except header) - gray background
        ('BACKGROUND', (1, 1), (1, -1), colors.white),  # Right column - white
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row bold
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),  # Left column bold
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(Paragraph("<i>How would the emergency affect you?</i>", ParagraphStyle('ItalicHeading', parent=normal_style, fontSize=11, textColor=colors.black)))
    story.append(emergency_affect_table)
    story.append(Spacer(1, 0.2*inch))
    
    # Section 3 heading
    story.append(Paragraph("3. My Emergency & Disaster Plan", heading_style))
    story.append(Paragraph("<i>Complete all applicable sections & if not applicable, mark as \"N/A\".</i>", normal_style))
    
    # Complete all applicable sections table
    # Define content for each section - bullet points go in LEFT column, field text goes in RIGHT column
    # Create a style for bold section names
    bold_text_style = ParagraphStyle(
        'BoldText',
        parent=table_text_style,
        fontSize=11,
        fontName='Helvetica-Bold',
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=12,
        leftIndent=0
    )
    
    communication_left = (
        "<b>Communication</b><br/><br/>"
        "[ ] I have my phone, computer, or tablet to be able to stay in touch with people or call people in an emergency.<br/><br/>"
        "[ ] I have informed my supports about the best way to communicate with me.<br/><br/>"
        "[ ] I have friends or family who maintain regular contact who will seek assistance if unable to contact me."
    )
    communication_right = "Other important information about my communication:\n\n\n"
    
    health_left = (
        "<b>Management of Health</b><br/><br/>"
        "[ ] I know if I'm in an emergency - call 000.<br/><br/>"
        "[ ] I have copies of concession cards, health insurance cards and prescriptions.<br/><br/>"
        "[ ] I have discussed with my doctor how I will access controlled medications during and after an emergency.<br/><br/>"
        "[ ] I registered for MyGov."
    )
    health_right = (
        "Instructions for people in my support network so they can help me collect what I need if I have to evacuate:\n\n\n"
        "Things I need to manage my health & medical devices:\n\n\n"
    )
    
    at_left = (
        "<b>Assistive Technology (AT)</b><br/><br/>"
        "[ ] I have a list of items I would need to take with me if I needed to leave my home."
    )
    at_right = "How I will transport critical equipment I have to evacuate:\n\n\n"
    
    support_left = (
        "<b>Personal Support</b><br/><br/>"
        "[ ] I have a plan for when I get separated from the people who normally provide assistance.<br/><br/>"
        "[ ] I have discussed my plan with my emergency contact."
    )
    support_right = "Write down the back-up plan for assistance in emergencies:\n\n\n"
    
    pets_left = (
        "<b>Assistance animals and pets</b><br/><br/>"
        "[ ] I have a plan for who will look after my animal in case of an emergency."
    )
    pets_right = "Write down your animals needs here:\n\n\n"
    
    transport_left = (
        "<b>Transportation</b><br/><br/>"
        "[ ] I have thought about different plans to make sure that we leave in time for safe evacuation."
    )
    transport_right = "\n\n\n"
    
    living_left = (
        "<b>Living Situation</b><br/><br/>"
        "[ ] My smoke alarms are tested regularly.<br/><br/>"
        "[ ] I have a fire extinguisher and/or fire blanket present.<br/><br/>"
        "[ ] I keep a mobility device (if applicable) by my bed in case I have to evacuate quickly."
    )
    living_right = (
        "<i>Contact Fire and Rescue Services in your state to see if you are eligible for a home safety visit.</i><br/><br/>"
        "Write any notes here:\n\n\n"
    )
    
    social_left = (
        "<b>Social Connectedness</b><br/><br/>"
        "[ ] I have a plan for staying connected and in touch with people.<br/><br/>"
        "[ ] I have introduced myself to my neighbours."
    )
    social_right = "Write any notes here:\n\n\n"
    
    sections_data = [
        [Paragraph('<b>My Emergency & Disaster Plan</b>', ParagraphStyle('TableTextBlue', parent=table_text_style, fontSize=11, textColor=colors.white)), ''],
        [Paragraph(communication_left, table_text_style), Paragraph(communication_right, table_text_style)],
        [Paragraph(health_left, table_text_style), Paragraph(health_right, table_text_style)],
        [Paragraph(at_left, table_text_style), Paragraph(at_right, table_text_style)],
        [Paragraph(support_left, table_text_style), Paragraph(support_right, table_text_style)],
        [Paragraph(pets_left, table_text_style), Paragraph(pets_right, table_text_style)],
        [Paragraph(transport_left, table_text_style), Paragraph(transport_right, table_text_style)],
        [Paragraph(living_left, table_text_style), Paragraph(living_right, table_text_style)],
        [Paragraph(social_left, table_text_style), Paragraph(social_right, table_text_style)],
        [Paragraph("<b>Other</b>", table_text_style), Paragraph("Write any notes here:\n\n\n\n", table_text_style)]
    ]
    
    sections_table = Table(sections_data, colWidths=[2.5*inch, 3.5*inch])
    # Define gray color for backgrounds
    GRAY_COLOR = colors.HexColor('#d9d9d9')
    
    sections_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), BLUE_COLOR),  # First row first column - blue
        ('TEXTCOLOR', (0, 0), (0, 0), colors.white),
        ('BACKGROUND', (1, 0), (1, 0), colors.white),  # First row second column
        ('BACKGROUND', (0, 1), (0, -1), GRAY_COLOR),  # Left column (except first row) - gray
        ('BACKGROUND', (1, 1), (1, -1), colors.white),  # Right column - white
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),  # All left column bold
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(sections_table)
    story.append(Spacer(1, 0.2*inch))
    
    # Final table with signatures (4 columns, 3 rows)
    # Get modification date - try to get from PDF file modification time or use current date
    from datetime import datetime
    mod_date = datetime.now().strftime('%d/%m/%Y')
    
    final_data = [
        [Paragraph("<b>Client's Name</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white)), Paragraph(client_name, table_text_style), Paragraph("<b>Team member's name</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white)), Paragraph(team_member_name, table_text_style)],
        [Paragraph("<b>Signature</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white)), '', Paragraph("<b>Signature</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white)), ''],
        [Paragraph("<b>Date</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white)), mod_date, Paragraph("<b>Date</b>", ParagraphStyle('TableHeader', parent=table_text_style, fontSize=11, textColor=colors.white)), mod_date]
    ]
    
    final_table = Table(final_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
    final_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), BLUE_COLOR),  # Client's Name
        ('BACKGROUND', (2, 0), (2, 0), BLUE_COLOR),  # Team member's name
        ('BACKGROUND', (0, 1), (0, 1), BLUE_COLOR),  # Signature (left)
        ('BACKGROUND', (2, 1), (2, 1), BLUE_COLOR),  # Signature (right)
        ('BACKGROUND', (0, 2), (0, 2), BLUE_COLOR),  # Date (left)
        ('BACKGROUND', (2, 2), (2, 2), BLUE_COLOR),  # Date (right)
        ('TEXTCOLOR', (0, 0), (0, 0), colors.white),
        ('TEXTCOLOR', (2, 0), (2, 0), colors.white),
        ('TEXTCOLOR', (0, 1), (0, 1), colors.white),
        ('TEXTCOLOR', (2, 1), (2, 1), colors.white),
        ('TEXTCOLOR', (0, 2), (0, 2), colors.white),
        ('TEXTCOLOR', (2, 2), (2, 2), colors.white),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),  # Data columns - white
        ('BACKGROUND', (3, 0), (3, -1), colors.white),
        ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
        ('TEXTCOLOR', (3, 0), (3, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('ALIGN', (2, 0), (2, -1), 'LEFT'),
        ('ALIGN', (3, 0), (3, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTNAME', (3, 0), (3, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(final_table)
    
    # Build PDF with headers and footers
    doc.build(story, onFirstPage=_add_first_page_header, onLaterPages=_add_header_footer)
    print("Emergency & Disaster Plan PDF created successfully!")

def extract_time_from_item_name(item_name):
    """
    Extract time information from support item name.
    Examples:
    - "Assistance With Self-Care Activities - Standard - Weekday Daytime" -> "Weekday Daytime"
    - "Assistance With Self-Care Activities - Standard - Weekday Night" -> "Weekday Night"
    - "Assistance With Self-Care Activities - Standard - Saturday" -> "Saturday"
    - "Assistance With Self-Care Activities - Standard - Public Holiday" -> "Public Holiday"
    """
    if not item_name:
        return ''
    
    # Common time patterns
    time_patterns = [
        'Weekday Daytime',
        'Weekday Night',
        'Weekday Evening',
        'Night-Time Sleepover',
        'Saturday',
        'Sunday',
        'Public Holiday',
        'Weekend',
        'After Hours',
        'Daytime',
        'Evening',
        'Night'
    ]
    
    # Check if any time pattern is in the name
    for pattern in time_patterns:
        if pattern in item_name:
            return pattern
    
    # If no pattern found, try to extract the last part after the last dash
    parts = item_name.split(' - ')
    if len(parts) > 1:
        last_part = parts[-1].strip()
        # If the last part looks like a time descriptor, return it
        if any(word in last_part for word in ['Day', 'Night', 'Evening', 'Weekend', 'Holiday', 'Saturday', 'Sunday']):
            return last_part
    
    return ''

def create_service_estimate_csv(csv_data, output_path, contact_name=None, ndis_items=None):
    """
    Create a Service Estimate CSV file from provided data dictionary.
    
    Args:
        csv_data: Dictionary containing form data
        output_path: Path where the CSV should be saved
        contact_name: Optional name (not used for CSV, but kept for consistency)
        ndis_items: Optional pre-loaded NDIS items (for performance)
    """
    # Load NDIS support items if not provided - need to also load the full CSV to get item names with time
    if ndis_items is None:
        ndis_items = load_ndis_support_items()
    
    # Also load the full NDIS CSV to get the actual item names (which contain time info)
    ndis_item_names = {}
    try:
        with open('outputs/other/NDIS Support Items - NDIS Support Items.csv', 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            for row in reader:
                item_name = row['Support Item Name'].strip()
                # Store mapping from normalized name to full name
                normalized_name = item_name.lower().strip()
                ndis_item_names[normalized_name] = item_name
    except Exception as e:
        print(f"Error loading NDIS item names: {e}")
    
    # Get team value to determine which state's price to use
    team_value = csv_data.get('Neighbourhood Care representative team', '')
    # Clean up checkbox characters
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Determine which state's price to use
    price_state = get_price_state(team_value)
    price_key = 'wa_price' if price_state == 'WA' else 'qld_price'
    
    # Extract support items 1-8 from Support Items Required section
    support_items_data = []
    for i in range(1, 9):  # Support items 1-8
        key = f'Support item ({i}) (Support Items Required)'
        item_name = csv_data.get(key, '').strip()
        
        if item_name:
            # Look up the item in NDIS items
            item_details = lookup_support_item(ndis_items, item_name)
            
            # Find the matching NDIS item name (which has time info)
            matched_ndis_name = item_name
            item_found = False
            if item_name in ndis_items:
                matched_ndis_name = item_name
                item_found = True
            else:
                # Try to find a match in ndis_item_names
                item_name_lower = item_name.lower().strip()
                for ndis_name_lower, ndis_name_full in ndis_item_names.items():
                    if item_name_lower in ndis_name_lower or ndis_name_lower in item_name_lower:
                        matched_ndis_name = ndis_name_full
                        item_found = True
                        break
            
            # Extract time from the matched NDIS item name (which has the time info)
            time = extract_time_from_item_name(matched_ndis_name)
            
            # Get the price based on state
            price = item_details.get(price_key, '')
            
            if item_found:
                support_items_data.append({
                    'Name': item_name,
                    'Category': 'Core',
                    'Number': item_details.get('number', ''),
                    'Unit': item_details.get('unit', ''),
                    'Price': price,
                    'Variable': 'TRUE',
                    'Time': time
                })
            else:
                # Item not found - still add it but with placeholder values
                # Try to extract time from the original item name as fallback
                time = extract_time_from_item_name(item_name)
                support_items_data.append({
                    'Name': item_name,
                    'Category': 'Core',
                    'Number': '[Not Found]',
                    'Unit': '[Not Found]',
                    'Price': '[Not Found]',
                    'Variable': 'TRUE',
                    'Time': time
                })
    
    # Column order: Name, Category, Number, Unit, Price, Variable, Time
    fieldnames = ['Name', 'Category', 'Number', 'Unit', 'Price', 'Variable', 'Time']
    
    # Write CSV file
    if support_items_data:
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(support_items_data)
        print(f"Service Estimate CSV created successfully with {len(support_items_data)} items!")
    else:
        # Create empty CSV with headers if no items found
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
        print("Service Estimate CSV created successfully (empty - no support items found)!")

def create_risk_assessment_from_data(csv_data, output_path, contact_name=None, active_users=None):
    """
    Create a Risk Assessment PDF from provided data dictionary.
    
    Args:
        csv_data: Dictionary containing form data
        output_path: Path where the PDF should be saved
        contact_name: Optional name for "Person Completing this assessment"
        active_users: Optional pre-loaded active users (for performance, not currently used but kept for consistency)
    """
    from datetime import datetime
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=0,
        leftIndent=0,
        fontName=get_calibri_bold_font()
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=11,
        textColor=BLUE_COLOR,
        alignment=TA_LEFT,
        spaceAfter=8,
        spaceBefore=12,
        leftIndent=0,
        fontName=get_calibri_font()
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=6,
        leading=14,
        leftIndent=0,
        fontName=get_calibri_font()
    )
    
    table_text_style = ParagraphStyle(
        'TableText',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceAfter=0,
        leading=12,
        leftIndent=0,
        fontName=get_calibri_font()
    )
    
    # Title
    story.append(Paragraph("PRO025 - Client Risk Assessment", title_style))
    
    # First table: Participant, Person Completing, Role, Date
    first_name = csv_data.get('First name (Details of the Client)', '').strip()
    surname = csv_data.get('Surname (Details of the Client)', '').strip()
    participant_name = ' '.join([p for p in [first_name, surname] if p]).strip() or ''
    person_completing = contact_name or ''
    role = 'Support Worker'
    assessment_date = datetime.now().strftime('%d/%m/%Y')
    
    # Create style for white text labels (bold)
    white_label_style = ParagraphStyle(
        'WhiteLabel',
        parent=table_text_style,
        textColor=colors.white,
        fontName=get_calibri_bold_font()
    )
    
    # Create style for table headers (bold)
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=table_text_style,
        fontSize=11,
        textColor=colors.white,
        alignment=TA_CENTER,
        fontName=get_calibri_bold_font()
    )
    
    first_table_data = [
        [Paragraph('Participant', white_label_style), Paragraph(participant_name + '\n\n', table_text_style)],
        [Paragraph('Person Completing this assessment', white_label_style), Paragraph(person_completing + '\n\n', table_text_style)],
        [Paragraph('Role', white_label_style), Paragraph(role + '\n\n', table_text_style)],
        [Paragraph('Date of Assessment', white_label_style), Paragraph(assessment_date + '\n\n', table_text_style)]
    ]
    
    # Define custom colors for Risk Assessment
    FIRST_TABLE_HEADER_COLOR = colors.HexColor('#027bc4')
    FIRST_TABLE_VALUE_COLOR = colors.HexColor('#d3dfee')
    ENV_HAZARDS_COLOR = colors.HexColor('#00b050')
    MEDICATION_COLOR = colors.HexColor('#943734')
    LIVING_ALONE_COLOR = colors.HexColor('#5f497a')
    LIVING_ALONE_FIELD_COLOR = colors.HexColor('#d3dfee')
    
    first_table = Table(first_table_data, colWidths=[2.8*inch, 3.2*inch])
    first_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), FIRST_TABLE_HEADER_COLOR),  # Left column (labels)
        ('BACKGROUND', (1, 0), (1, -1), FIRST_TABLE_VALUE_COLOR),  # Right column (values)
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),  # White text for labels
        ('TEXTCOLOR', (1, 0), (1, -1), colors.black),  # Black text for values
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), get_calibri_bold_font()),
        ('FONTNAME', (1, 0), (1, -1), get_calibri_font()),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(first_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Common Hazards table
    common_hazards = [
        'Lifting, supporting and transferring',
        'Personal care (e.g. manual handling, slip trips and falls, biological hazards, humidity etc during showering, sponging and toileting)',
        'Using equipment or assistive technology',
        'Client behaviours of concern that may harm themselves (is there a Behaviour Support Plan (BSP) in place?)',
        'Client behaviours of concern that may harm others (BSP)?',
        'Other people\'s behaviours of concern that may harm client of worker',
        'Are there any restrictive practices in place? (BSP?)',
        'Client\'s ability to communicate needs and wants clearly, including in emergencies. Is English their second language?',
        'Limited informal supports',
        'Limited formal supports or access to formal supports',
        'Reduced choice and control',
        'Reduced opportunities for independence',
        'Social isolation',
        'Dignity',
        'Self esteem',
        'Health conditions and wellbeing (e.g. epilepsy, heart condition etc)',
        'Other'
    ]
    
    common_hazards_data = [
        [Paragraph('Common Hazards', table_header_style),
         Paragraph('Risk Identified', table_header_style),
         Paragraph('Who is at Risk', table_header_style),
         Paragraph('Proposed Control (Action taken to prevent/minimise the risk)', table_header_style),
         Paragraph('Implemented (Date)', table_header_style)]
    ]
    
    for hazard in common_hazards:
        common_hazards_data.append([
            Paragraph(hazard + '\n\n', table_text_style),
            '\n\n',
            '\n\n',
            '\n\n',
            '\n\n'
        ])
    
    common_hazards_table = Table(common_hazards_data, colWidths=[1.8*inch, 1.2*inch, 1.2*inch, 1.8*inch, 1*inch])
    common_hazards_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FIRST_TABLE_HEADER_COLOR),  # Header row - #027bc4
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), CALIBRI_BOLD_FONT),
        ('FONTNAME', (0, 1), (-1, -1), CALIBRI_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(common_hazards_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Environmental Hazards table
    environmental_hazards = [
        'Slips, trips and falls risks',
        'Air quality (including cigarette smoke)',
        'Pets or other animals that could be dangerous to client, worker or others',
        'High crime area',
        'Hazardous materials/liquids',
        'Hot water',
        'Sharp objects',
        'Temperature',
        'Are there smoke alarms?',
        'Are there fire extinguishers?',
        'Are there RCDs (safety fuses)?',
        'Stairs',
        'Lighting',
        'Is there mobile phone reception and/or working landline?',
        'Are there any other occupants or visitors likely to be present during home visits?',
        'Other (e.g. weapons, firearms)',
        'Other'
    ]
    
    environmental_hazards_data = [
        [Paragraph('Environmental Hazards', table_header_style),
         Paragraph('Risk Identified', table_header_style),
         Paragraph('Who is at Risk', table_header_style),
         Paragraph('Proposed Control (Action taken to prevent/minimise the risk)', table_header_style),
         Paragraph('Implemented (Date)', table_header_style)]
    ]
    
    for hazard in environmental_hazards:
        environmental_hazards_data.append([
            Paragraph(hazard + '\n\n', table_text_style),
            '\n\n',
            '\n\n',
            '\n\n',
            '\n\n'
        ])
    
    environmental_hazards_table = Table(environmental_hazards_data, colWidths=[1.8*inch, 1.2*inch, 1.2*inch, 1.8*inch, 1*inch])
    environmental_hazards_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), ENV_HAZARDS_COLOR),  # Header row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), CALIBRI_BOLD_FONT),
        ('FONTNAME', (0, 1), (-1, -1), CALIBRI_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(environmental_hazards_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Medication table
    medication_data = [
        [Paragraph('Medication', table_header_style),
         Paragraph('Dose and time', table_header_style),
         Paragraph('Side effects or interactions to be aware of', table_header_style),
         Paragraph('Administered by', table_header_style),
         Paragraph('Is Neighbourhood Care involved in medication management?', table_header_style)]
    ]
    
    # Add empty rows for medication entries
    for i in range(5):
        medication_data.append(['\n\n', '\n\n', '\n\n', '\n\n', '\n\n'])
    
    medication_table = Table(medication_data, colWidths=[1.5*inch, 1.2*inch, 1.5*inch, 1.2*inch, 1.6*inch])
    medication_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), MEDICATION_COLOR),  # Header row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), CALIBRI_BOLD_FONT),
        ('FONTNAME', (0, 1), (-1, -1), CALIBRI_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(medication_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Living Alone Assessment table
    living_alone_fields = [
        'Do you feel safe at home, work, in the community?',
        'Are there any places you don\'t feel safe?',
        'Do you have any plans in place for emergency situations (e.g. fire, illness, injury, severe weather)?',
        'Do you have safety equipment at home (e.g. fire extinguisher, fire blanket)?',
        'Do you feel safe at home, work, in the community?',
        'Do you have any plans in place for emergency situations (e.g. fire, illness, injury, severe weather)?',
        'Have you ever been injured, fallen etc?',
        'Informal supports: family and friends - who visits and how often?',
        'Paid supports - which services are engaged, number of services and support workers, number of visitors per week estimate',
        'What is your ability to contact people if you need their help (e.g. family, friends, staff by phone or by other method)?'
    ]
    
    living_alone_data = [
        # Title row - merged across all columns
        [Paragraph('Living Alone Assessment (To be completed if the client is living alone)', table_header_style),
         '', '', ''],
        # Header row
        [Paragraph('Identified Gap / Risk / Hazard', table_header_style),
         Paragraph('Notes', table_header_style),
         Paragraph('Proposed Control (Action taken to prevent/minimise the risk)', table_header_style),
         Paragraph('Implemented (Date)', table_header_style)]
    ]
    
    for field in living_alone_fields:
        living_alone_data.append([
            Paragraph(field + '\n\n', table_text_style),
            '\n\n',
            '\n\n',
            '\n\n'
        ])
    
    # Create bold style for Violence/Abuse section
    bold_text_style = ParagraphStyle(
        'BoldText',
        parent=table_text_style,
        fontName=get_calibri_bold_font()
    )
    
    # Add Violence, Abuse section - span all columns
    living_alone_data.append([
        Paragraph('Violence, Abuse, Sexual Abuse, Discrimination, Exploitation', bold_text_style),
        '', '', ''
    ])
    
    living_alone_data.append([
        Paragraph('Are there any people you do not feel safe with?* Can you tell me more about this?\n\n', table_text_style),
        '\n\n',
        '\n\n',
        '\n\n'
    ])
    
    # Add note - span all columns
    living_alone_data.append([
        Paragraph('(*Only ask questions below if triggered by a positive response to above question.)', bold_text_style),
        '', '', ''
    ])
    
    violence_abuse_fields = [
        'Have you ever been hurt by anyone?',
        'Has anyone ever taken advantage of you?',
        'Does anyone yell or curse at you?',
        'If yes to any of the above questions: if so, who did you tell?',
        'What was done to address your concerns?'
    ]
    
    for field in violence_abuse_fields:
        living_alone_data.append([
            Paragraph(field + '\n\n', table_text_style),
            '\n\n',
            '\n\n',
            '\n\n'
        ])
    
    # Find rows for Violence/Abuse and note to span across all columns
    # Violence/Abuse is after the 10 living_alone_fields, so row index (0-based: title=0, header=1, fields=2-11, violence=12)
    violence_row = 2 + len(living_alone_fields)  # Row index for Violence/Abuse
    note_row = violence_row + 2  # Row index for the note
    
    living_alone_table = Table(living_alone_data, colWidths=[2*inch, 1.5*inch, 2*inch, 1.5*inch])
    living_alone_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), LIVING_ALONE_COLOR),  # Title row
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('SPAN', (0, 0), (-1, 0)),  # Merge title row across all columns
        ('BACKGROUND', (0, 1), (-1, 1), LIVING_ALONE_COLOR),  # Header row
        ('TEXTCOLOR', (0, 1), (-1, 1), colors.white),
        ('BACKGROUND', (0, 2), (-1, -1), LIVING_ALONE_FIELD_COLOR),  # All data rows
        ('TEXTCOLOR', (0, 2), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (0, 1), (-1, 1), 'CENTER'),
        ('ALIGN', (0, 2), (0, -1), 'LEFT'),
        ('ALIGN', (1, 2), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 1), CALIBRI_BOLD_FONT),
        ('FONTNAME', (0, 2), (-1, -1), CALIBRI_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('SPAN', (0, violence_row), (-1, violence_row)),  # Merge Violence/Abuse row
        ('SPAN', (0, note_row), (-1, note_row))  # Merge note row
    ]))
    
    story.append(living_alone_table)
    
    # Build PDF with headers and footers
    doc.build(story, onFirstPage=_add_first_page_header, onLaterPages=_add_header_footer)
    print("Risk Assessment PDF created successfully!")

def create_support_plan_from_data(csv_data, output_path, contact_name=None, active_users=None):
    """
    Create a Support Plan Word document (.docx) from provided data dictionary.
    
    Args:
        csv_data: Dictionary containing form data
        output_path: Path where the .docx should be saved
        contact_name: Optional name to use for Key Contact lookup
        active_users: Optional pre-loaded active users (for performance)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from datetime import datetime
        import re
    except ImportError:
        raise ImportError("python-docx is required for Support Plan generation. Please install it: pip install python-docx")
    
    # Get team value to determine which active users CSV to use
    team_value = csv_data.get('Neighbourhood Care representative team', '')
    # Clean up checkbox characters
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    
    # Load active users based on team if not provided
    if active_users is None:
        active_users = load_active_users(team_value)
    
    # Get key contact information
    key_contact_name_to_use = contact_name or csv_data.get('Respondent', '')
    key_contact_data = lookup_user_data(active_users, key_contact_name_to_use) if key_contact_name_to_use else {'name': '', 'mobile': '', 'email': '', 'team': ''}
    
    # Extract client information
    first_name = csv_data.get('First name (Details of the Client)', '').strip()
    surname = csv_data.get('Surname (Details of the Client)', '').strip()
    dob_str = csv_data.get('Date of birth (Details of the Client)', '').strip()
    home_address = csv_data.get('Home address (Contact Details of the Client)', '').strip()
    
    # Create Word document
    doc = Document()
    
    # Set default font to Calibri, size 12, centered
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header and footer
    section = doc.sections[0]
    
    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Try to add image to header
    image_filename = 'image.png'
    script_dir = os.path.dirname(os.path.abspath(__file__))
    image_path = None
    search_dirs = [script_dir, os.getcwd(), '.']
    
    for search_dir in search_dirs:
        if os.path.exists(search_dir):
            try:
                test_path = os.path.join(search_dir, image_filename)
                if os.path.exists(test_path):
                    image_path = os.path.abspath(test_path)
                    break
                for filename in os.listdir(search_dir):
                    if filename.lower() == image_filename.lower() or (filename.lower().startswith('image') and filename.lower().endswith('.png')):
                        full_path = os.path.join(search_dir, filename)
                        if os.path.exists(full_path):
                            image_path = os.path.abspath(full_path)
                            break
                if image_path:
                    break
            except Exception:
                continue
    
    if image_path and os.path.exists(image_path):
        try:
            from docx.shared import Inches
            run = header_para.add_run()
            run.add_picture(image_path, width=Inches(1.5))
        except Exception:
            pass
    
    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Neighbourhood Care | Suite 103, 19 Ogilvie Road, Mount Pleasant, WA 6153 | ABN 40 634 832 607")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)  # #7F7F7F
    
    # Add page number to footer (right side)
    footer_para.add_run("  ")
    page_num_run = footer_para.add_run()
    page_num_run._element.text = ""
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_num_run._element.append(fldChar1)
    page_num_run._element.append(instrText)
    page_num_run._element.append(fldChar2)
    
    # Define the color for text and borders
    border_color = RGBColor(0x25, 0x6e, 0xb7)  # #256eb7
    
    # Helper function to create a boxed section
    def create_boxed_section():
        """Create a table with one cell that acts as a box"""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        
        # Set cell padding
        tc_pr = cell._element.get_or_add_tcPr()
        tc_mar = OxmlElement('w:tcMar')
        for margin in ['top', 'left', 'bottom', 'right']:
            margin_elem = OxmlElement(f'w:{margin}')
            margin_elem.set(qn('w:w'), '144')  # 0.1 inch
            margin_elem.set(qn('w:type'), 'dxa')
            tc_mar.append(margin_elem)
        tc_pr.append(tc_mar)
        
        # Set border color to #256eb7
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '256EB7')  # #256eb7
            tc_borders.append(border)
        tc_pr.append(tc_borders)
        
        return cell
    
    # Helper function to add paragraph with no spacing
    def add_paragraph_no_spacing(cell, text=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Add a paragraph with no space before or after"""
        if text:
            p = cell.add_paragraph(text)
        else:
            p = cell.add_paragraph()
        p.alignment = alignment
        # Force zero spacing using XML to override any style defaults
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0  # Single line spacing
        # Also set via XML to ensure it's truly zero
        pPr = p._element.get_or_add_pPr()
        # Remove any existing spacing
        for spacing_elem in pPr.xpath('.//w:spacing'):
            pPr.remove(spacing_elem)
        # Add explicit zero spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Single line spacing (240 twips = 12pt)
        spacing.set(qn('w:lineRule'), 'exact')  # Use exact line spacing instead of auto
        pPr.append(spacing)
        return p
    
    # Helper function to ensure font size 12 for runs
    def set_font_size_12(run):
        """Set font size to 12 for a run"""
        run.font.size = Pt(12)
    
    # Helper function to add an empty paragraph that's actually visible (for spacing)
    def add_empty_line(cell):
        """Add an empty paragraph that will be visible as a blank line"""
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set minimal spacing to ensure the empty line is visible
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # Set line height via XML to ensure it's visible
        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Single line spacing (240 twips = 12pt)
        spacing.set(qn('w:lineRule'), 'exact')  # Use exact line spacing instead of auto
        pPr.append(spacing)
        return p
    
    # Title box - "My Support Plan"
    title_cell = create_boxed_section()
    # Fill background with #256eb7
    tc_pr = title_cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '256EB7')  # #256eb7
    tc_pr.append(shd)
    
    p = add_paragraph_no_spacing(title_cell)
    run = p.add_run('My Support Plan')
    run.font.color.rgb = RGBColor(255, 255, 255)  # White
    run.bold = True
    run.font.size = Pt(18)
    
    doc.add_paragraph()  # Empty line after title
    
    # Header section - no spaces between paragraphs
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run1 = p.add_run('My Name: ')
    run1.font.color.rgb = border_color
    run1.font.size = Pt(12)
    run2 = p.add_run(f'{first_name} {surname}'.strip() if (first_name or surname) else '')
    run2.font.color.rgb = border_color
    run2.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run1 = p.add_run('My Date of Birth: ')
    run1.font.color.rgb = border_color
    run1.font.size = Pt(12)
    run2 = p.add_run(dob_str if dob_str else '')
    run2.font.color.rgb = border_color
    run2.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    run1 = p.add_run('My Address: ')
    run1.font.color.rgb = border_color
    run1.font.size = Pt(12)
    run2 = p.add_run(home_address if home_address else '')
    run2.font.color.rgb = border_color
    run2.font.size = Pt(12)
    
    doc.add_paragraph()  # One empty line between "My Address:" and "About this Plan" box
    
    # About this Plan section - in one box
    about_plan_cell = create_boxed_section()
    p = add_paragraph_no_spacing(about_plan_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('About this Plan')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    
    bullet_points = [
        'This plan lets you share information about who you are, what your life is like and your dreams',
        'You can make this plan by yourself, with your support worker or with someone you choose',
        'This plan contains your goals and what supports you need to help you achieve them',
        'This plan has the supports you have now around you and how they can help you achieve your goals'
    ]
    for i, point in enumerate(bullet_points):
        p = about_plan_cell.add_paragraph(point, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)  # Remove space below all bullet points
        p.paragraph_format.line_spacing = 1.0  # Single line spacing
        # Force zero spacing via XML to override style defaults
        pPr = p._element.get_or_add_pPr()
        existing_spacing = pPr.xpath('.//w:spacing')
        for spacing_elem in existing_spacing:
            pPr.remove(spacing_elem)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Single line spacing (240 twips = 12pt)
        spacing.set(qn('w:lineRule'), 'exact')  # Use exact line spacing instead of auto
        pPr.append(spacing)
        for run in p.runs:
            set_font_size_12(run)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My Support Team section - in one box
    support_team_cell = create_boxed_section()
    p = add_paragraph_no_spacing(support_team_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run1 = p.add_run('My Support Team: ')  # Not bold
    set_font_size_12(run1)
    run2 = p.add_run(key_contact_data.get('team', '') if key_contact_data.get('team') else '')
    set_font_size_12(run2)
    
    p = add_paragraph_no_spacing(support_team_cell)
    run1 = p.add_run('My Key Contact: ')  # Not bold
    set_font_size_12(run1)
    run2 = p.add_run(key_contact_data.get('name', '') if key_contact_data.get('name') and key_contact_data.get('name') != '[Not Found]' else '')
    set_font_size_12(run2)
    
    p = add_paragraph_no_spacing(support_team_cell)
    run1 = p.add_run('Contact Number: ')  # Not bold
    set_font_size_12(run1)
    run2 = p.add_run(key_contact_data.get('mobile', '') if key_contact_data.get('mobile') and key_contact_data.get('mobile') != '[Not Found]' else '')
    set_font_size_12(run2)
    
    p = add_paragraph_no_spacing(support_team_cell)
    run1 = p.add_run('Email: ')  # Not bold
    set_font_size_12(run1)
    run2 = p.add_run(key_contact_data.get('email', '') if key_contact_data.get('email') and key_contact_data.get('email') != '[Not Found]' else '')
    set_font_size_12(run2)
    
    doc.add_paragraph()  # Empty line
    
    # What are some of the things section
    p = doc.add_paragraph('What are some of the things that you want the people supporting you to know about you?')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font_size_12(run)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # About Me box
    about_me_cell = create_boxed_section()
    p = add_paragraph_no_spacing(about_me_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('About Me')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(about_me_cell)
    run = p.add_run('For example, your living situation, study, friends, family/relationships, your personality, things that are important to you, how you spend your leisure time')
    run.italic = True
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(about_me_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My NDIS Goals box
    ndis_goals_cell = create_boxed_section()
    p = add_paragraph_no_spacing(ndis_goals_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('My NDIS Goals')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(ndis_goals_cell)
    run = p.add_run('Short term goals')
    run.font.color.rgb = border_color
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(ndis_goals_cell)
    p = add_paragraph_no_spacing(ndis_goals_cell)
    run = p.add_run('Medium & Long term goals')
    run.font.color.rgb = border_color
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(ndis_goals_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # Gift of the Head, Heart & Hand box
    gift_cell = create_boxed_section()
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('Gift of the Head, Heart & Hand')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    add_paragraph_no_spacing(gift_cell)  # Empty line after "Gift of the Head, Heart & Hand"
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('GIFTS OF THE HEAD')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('(What special knowledge, expertise, life experience do you have that you can share with others?)')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(gift_cell)
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('GIFTS OF THE HEART')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('(What things are really important to you, that you deeply care about and would welcome to share with others?)')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(gift_cell)
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('GIFTS OF THE HAND')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(gift_cell)
    run = p.add_run('(What practical skill do you bring with you, that you are good at, proud of and you may wish to share with others?)')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(gift_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My Dreams box
    dreams_cell = create_boxed_section()
    p = add_paragraph_no_spacing(dreams_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('My Dreams')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(dreams_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # People in My Life box
    people_cell = create_boxed_section()
    p = add_paragraph_no_spacing(people_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('People in My Life')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(people_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My Week box
    week_cell = create_boxed_section()
    p = add_paragraph_no_spacing(week_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('My Week')
    run.font.color.rgb = border_color
    run.bold = True
    add_paragraph_no_spacing(week_cell)  # Empty line after "My Week"
    p = add_paragraph_no_spacing(week_cell)
    p.paragraph_format.space_after = Pt(12)  # Add space after description text
    run = p.add_run('Identify when you currently have support with day to day activities and when you feel you need additional support. This might be from formal or informal supports')
    set_font_size_12(run)
    
    # Add table inside the box - centered with proper spacing
    # Add a paragraph before table for spacing
    p = add_paragraph_no_spacing(week_cell)
    p.paragraph_format.space_after = Pt(6)  # Small space before table
    
    week_table = week_cell.add_table(rows=6, cols=8)
    week_table.style = 'Table Grid'
    # Center the table by setting alignment on the table element
    tbl_pr = week_table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        week_table._element.insert(0, tbl_pr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tbl_pr.append(jc)
    
    # Set table width to be smaller so it's not squished and center it
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), '7200')  # 5 inches (smaller to prevent squishing)
    tbl_width.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_width)
    
    # Set table border color to #256eb7 for all cells
    for row in week_table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '256EB7')  # #256eb7
                tc_borders.append(border)
            tc_pr.append(tc_borders)
    
    # Header row
    days = ['', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    header_cells = week_table.rows[0].cells
    for i, day in enumerate(days):
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if day:  # Only bold the day names, not the empty first cell
            run = p.add_run(day)
            run.font.color.rgb = border_color
            run.bold = True
    
    # Time rows
    times = ['Early Morning', 'Morning', 'Afternoon', 'Evening', 'Overnight']
    for i, time in enumerate(times):
        p = week_table.rows[i + 1].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(time)
        run.font.color.rgb = border_color
        # Center align all other cells in this row
        for j in range(1, 8):
            week_table.rows[i + 1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My Safety box
    safety_cell = create_boxed_section()
    p = add_paragraph_no_spacing(safety_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    p.paragraph_format.space_after = Pt(0)  # Ensure no space below
    run = p.add_run('My Safety')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(safety_cell)
    run = p.add_run('Following on from the risk assessment, were there people, places or times that you feel unsafe? What changes need to be made and what support is needed so that you feel safe? Is there a formal safety plan in place? Is one needed?')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(safety_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My Medications box
    med_cell = create_boxed_section()
    p = add_paragraph_no_spacing(med_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    p.paragraph_format.space_after = Pt(0)  # Ensure no space below
    run = p.add_run('My Medications and how I manage them')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(med_cell)
    run = p.add_run('Do you need assistance with organising and taking your medication?')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(med_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My special supports box
    special_cell = create_boxed_section()
    p = add_paragraph_no_spacing(special_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    p.paragraph_format.space_after = Pt(0)  # Ensure no space below
    run = p.add_run('My special supports')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(special_cell)
    run = p.add_run('Do you have any special needs or equipment and do you have plans already to help make sure your support workers know how to care for you such as:')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(special_cell)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # My Goals box
    goals_cell = create_boxed_section()
    p = add_paragraph_no_spacing(goals_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('My Goals')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    add_paragraph_no_spacing(goals_cell)  # Empty line after "My Goals"
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('My SMART Goal 1')
    set_font_size_12(run)
    add_paragraph_no_spacing(goals_cell)  # Empty line
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('Strategies - What will help me achieve my goal? Who will help me achieve my goal? What supports will I need?')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(goals_cell)
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('My SMART Goal 2')
    set_font_size_12(run)
    add_paragraph_no_spacing(goals_cell)  # Empty line
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('Strategies - What will help me achieve my goal? Who will help me achieve my goal? What supports will I need?')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(goals_cell)
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('My SMART Goal 3')
    set_font_size_12(run)
    add_paragraph_no_spacing(goals_cell)  # Empty line
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('Strategies - What will help me achieve my goal? Who will help me achieve my goal? What supports will I need?')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(goals_cell)
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('My SMART Goal 4')
    set_font_size_12(run)
    p = add_paragraph_no_spacing(goals_cell)
    run = p.add_run('Strategies - What will help me achieve my goal? Who will help me achieve my goal? What supports will I need?')
    set_font_size_12(run)
    for _ in range(4):
        add_paragraph_no_spacing(goals_cell)
    
    # How I Will Celebrate box
    celebrate_cell = create_boxed_section()
    p = add_paragraph_no_spacing(celebrate_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    run = p.add_run('How I Will Celebrate Achieving My Goals')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    add_empty_line(celebrate_cell)  # Empty line after "How I Will Celebrate Achieving My Goals"
    p = add_paragraph_no_spacing(celebrate_cell)
    run = p.add_run('Goal 1')
    run.font.color.rgb = border_color
    set_font_size_12(run)
    add_empty_line(celebrate_cell)  # Empty line between goals
    p = add_paragraph_no_spacing(celebrate_cell)
    run = p.add_run('Goal 2')
    run.font.color.rgb = border_color
    set_font_size_12(run)
    add_empty_line(celebrate_cell)  # Empty line between goals
    p = add_paragraph_no_spacing(celebrate_cell)
    run = p.add_run('Goal 3')
    run.font.color.rgb = border_color
    set_font_size_12(run)
    add_empty_line(celebrate_cell)  # Empty line between goals
    p = add_paragraph_no_spacing(celebrate_cell)
    run = p.add_run('Goal 4')
    run.font.color.rgb = border_color
    set_font_size_12(run)
    
    doc.add_paragraph()  # Empty line between boxes
    
    # Final signature section - in a box
    signature_cell = create_boxed_section()
    p = add_paragraph_no_spacing(signature_cell)
    p.paragraph_format.space_before = Pt(0)  # Ensure no space above
    p.paragraph_format.space_after = Pt(12)  # One space below "This Is My Plan"
    run = p.add_run('This Is My Plan')
    run.font.color.rgb = border_color
    run.bold = True
    set_font_size_12(run)
    
    p = add_paragraph_no_spacing(signature_cell)
    run1 = p.add_run('Signature: ')
    run1.font.color.rgb = border_color
    run1.bold = True
    set_font_size_12(run1)
    run2 = p.add_run(f'{first_name} {surname}'.strip() if (first_name or surname) else '')
    run2.font.color.rgb = border_color  # Data following Signature: should be colored
    run2.bold = True  # Bold the data
    set_font_size_12(run2)
    
    add_paragraph_no_spacing(signature_cell)  # Empty line between Signature and Date
    
    p = add_paragraph_no_spacing(signature_cell)
    run1 = p.add_run('Date: ')
    run1.font.color.rgb = border_color
    run1.bold = True
    set_font_size_12(run1)
    run2 = p.add_run(datetime.now().strftime('%d/%m/%Y'))
    run2.font.color.rgb = border_color
    run2.bold = True  # Bold the data
    set_font_size_12(run2)
    
    # Save document
    doc.save(output_path)
    print(f"Support Plan Word document created successfully: {output_path}")

def create_medication_assistance_plan_from_data(csv_data, output_path, contact_name=None, active_users=None):
    """
    Create a Medication Assistance Plan DOCX document from CSV data
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from datetime import datetime
        import re
    except ImportError:
        raise ImportError("python-docx is required for Medication Assistance Plan generation. Please install it: pip install python-docx")
    
    # Extract client information
    first_name = csv_data.get('First name (Details of the Client)', '').strip()
    surname = csv_data.get('Surname (Details of the Client)', '').strip()
    dob_str = csv_data.get('Date of birth (Details of the Client)', '').strip()
    ndis_number = csv_data.get('NDIS number (Details of the Client)', '').strip()
    medicare_number = csv_data.get('Medicare number (Details of the Client)', '').strip() if csv_data.get('Medicare number (Details of the Client)') else ''
    
    # Get key contact name (similar to service agreement)
    team_value = csv_data.get('Neighbourhood Care representative team', '')
    team_value = team_value.replace('\uf0d7', '').replace('•', '').replace('●', '').replace('☐', '').replace('☑', '').replace('✓', '').strip()
    contact_name_to_use = contact_name or csv_data.get('Respondent', '')
    if active_users and contact_name_to_use:
        user_data = lookup_user_data(active_users, contact_name_to_use)
        key_contact_name = user_data.get('name', contact_name_to_use) if user_data else contact_name_to_use
    else:
        key_contact_name = contact_name_to_use
    
    # Extract assistance required fields
    def find_assistance_field(field_name):
        """Helper to find assistance-related fields in PDF data"""
        # Try exact match first
        value = csv_data.get(field_name, '').strip()
        if value:
            return value
        # Try variations
        field_lower = field_name.lower()
        for key, val in csv_data.items():
            if field_lower in key.lower() and val and str(val).strip():
                return str(val).strip()
        return ''
    
    communication_assistance = find_assistance_field('If applicable, describe the communication assistance required')
    medication_assistance_needed = csv_data.get('Does the client need assistance with their medication?', '').strip()
    equipment_assistive = find_assistance_field('Equipment or assistive technologies used')
    assisted_transfers = find_assistance_field('If applicable, describe the client\'s requirements for assisted transfers')
    catheter_management = find_assistance_field('If applicable, describe the client\'s requirements for catheter management')
    wound_pressure_care = find_assistance_field('If applicable, describe the client\'s requirements regarding wound care and/or pressure care')
    bowel_care = find_assistance_field('If applicable, describe the client\'s requirements regarding complex bowel care')
    enteral_feeding = find_assistance_field('If applicable, describe the client\'s requirements regarding enteral feeding and management')
    peg_feeding = find_assistance_field('If applicable, describe the client\'s requirements regarding PEG feeding')
    stoma_care = find_assistance_field('If applicable, describe the client\'s requirements regarding stoma care')
    additional_care = find_assistance_field('Additional care requirements')
    behaviour_support = find_assistance_field('Behaviour support requirements')
    
    # Create Word document
    doc = Document()
    
    # Set default font to Calibri, size 12, left-aligned
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add header and footer (same as Support Plan)
    section = doc.sections[0]
    
    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Try to add image to header
    image_filename = 'image.png'
    script_dir = os.path.dirname(os.path.abspath(__file__))
    image_path = None
    search_dirs = [script_dir, os.getcwd(), '.']
    
    for search_dir in search_dirs:
        if os.path.exists(search_dir):
            try:
                test_path = os.path.join(search_dir, image_filename)
                if os.path.exists(test_path):
                    image_path = os.path.abspath(test_path)
                    break
                for filename in os.listdir(search_dir):
                    if filename.lower() == image_filename.lower() or (filename.lower().startswith('image') and filename.lower().endswith('.png')):
                        full_path = os.path.join(search_dir, filename)
                        if os.path.exists(full_path):
                            image_path = os.path.abspath(full_path)
                            break
                if image_path:
                    break
            except Exception:
                continue
    
    if image_path and os.path.exists(image_path):
        try:
            from docx.shared import Inches
            run = header_para.add_run()
            run.add_picture(image_path, width=Inches(1.5))
        except Exception:
            pass
    
    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Neighbourhood Care | Suite 103, 19 Ogilvie Road, Mount Pleasant, WA 6153 | ABN 40 634 832 607")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)  # #7F7F7F
    
    # Add page number to footer (right side)
    footer_para.add_run("  ")
    page_num_run = footer_para.add_run()
    page_num_run._element.text = ""
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_num_run._element.append(fldChar1)
    page_num_run._element.append(instrText)
    page_num_run._element.append(fldChar2)
    
    # Define the color for text
    text_color = RGBColor(0x00, 0x7b, 0xc4)  # #007bc4
    box_fill_color = RGBColor(0xe0, 0xf4, 0xff)  # #e0f4ff
    
    # Helper function to create a boxed section
    def create_boxed_section():
        """Create a table with one cell that acts as a box with no outline but filled background"""
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        
        # Set cell padding
        tc_pr = cell._element.get_or_add_tcPr()
        tc_mar = OxmlElement('w:tcMar')
        for margin in ['top', 'left', 'bottom', 'right']:
            margin_elem = OxmlElement(f'w:{margin}')
            margin_elem.set(qn('w:w'), '144')  # 0.1 inch
            margin_elem.set(qn('w:type'), 'dxa')
            tc_mar.append(margin_elem)
        tc_pr.append(tc_mar)
        
        # Remove all borders (no outline)
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')  # No border
            tc_borders.append(border)
        tc_pr.append(tc_borders)
        
        # Set fill color to #e0f4ff
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E0F4FF')  # #e0f4ff
        shd.set(qn('w:val'), 'clear')
        tc_pr.append(shd)
        
        return cell
    
    # Helper function to add paragraph with no spacing
    def add_paragraph_no_spacing(cell, text=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Add a paragraph with no space before or after"""
        if text:
            p = cell.add_paragraph(text)
        else:
            p = cell.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # Also set via XML to ensure it's truly zero
        pPr = p._element.get_or_add_pPr()
        existing_spacing = pPr.xpath('.//w:spacing')
        for spacing_elem in existing_spacing:
            pPr.remove(spacing_elem)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # Single line spacing (240 twips = 12pt)
        spacing.set(qn('w:lineRule'), 'exact')  # Use exact line spacing instead of auto
        pPr.append(spacing)
        return p
    
    # Helper function to remove ALL spacing from nested table cells (more aggressive)
    def remove_all_spacing_from_nested_cell(cell):
        """Remove ALL spacing from nested table cells - maximum compactness"""
        # Remove cell margins/padding - set to absolute zero
        tc_pr = cell._element.get_or_add_tcPr()
        # Remove existing margins
        for margin_elem in tc_pr.xpath('.//w:tcMar'):
            tc_pr.remove(margin_elem)
        # Set all margins to absolute zero
        tc_mar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin_elem = OxmlElement(f'w:{margin_name}')
            margin_elem.set(qn('w:w'), '0')
            margin_elem.set(qn('w:type'), 'dxa')
            tc_mar.append(margin_elem)
        tc_pr.append(tc_mar)
        
        # Remove spacing from all paragraphs - absolute zero
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Modify XML directly
            pPr = paragraph._element.get_or_add_pPr()
            # Remove all spacing elements
            for spacing_elem in pPr.xpath('.//w:spacing'):
                pPr.remove(spacing_elem)
            
            # Add zero spacing with minimal line spacing
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            spacing.set(qn('w:line'), '100')  # Very tight line spacing (100 twips = ~5pt)
            spacing.set(qn('w:lineRule'), 'exact')
            pPr.append(spacing)
    
    # Helper function to remove all spacing from all paragraphs in a cell
    # Note: Google Docs may interpret spacing differently than Word
    def remove_all_spacing_from_cell(cell):
        """Remove spacing from all paragraphs in a cell and remove cell margins"""
        # Remove cell margins/padding - set to minimal value for Google Docs compatibility
        tc_pr = cell._element.get_or_add_tcPr()
        # Remove existing margins
        for margin_elem in tc_pr.xpath('.//w:tcMar'):
            tc_pr.remove(margin_elem)
        # Set all margins to zero for maximum compactness
        tc_mar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin_elem = OxmlElement(f'w:{margin_name}')
            margin_elem.set(qn('w:w'), '0')  # Zero margins for maximum compactness
            margin_elem.set(qn('w:type'), 'dxa')
            tc_mar.append(margin_elem)
        tc_pr.append(tc_mar)
        
        # Remove spacing from all paragraphs - use zero values for maximum compactness
        for paragraph in cell.paragraphs:
            # Set paragraph format first
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Then modify XML directly
            pPr = paragraph._element.get_or_add_pPr()
            # Remove all spacing elements
            for spacing_elem in pPr.xpath('.//w:spacing'):
                pPr.remove(spacing_elem)
            
            # Add zero spacing with tighter line spacing for maximum compactness
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')  # Zero spacing
            spacing.set(qn('w:after'), '0')  # Zero spacing
            spacing.set(qn('w:line'), '120')  # Tighter line spacing (120 twips = 6pt, half of normal)
            spacing.set(qn('w:lineRule'), 'exact')  # Use exact line spacing
            pPr.append(spacing)
            
            # Also try setting keepNext and keepLines to prevent extra spacing
            keepNext = OxmlElement('w:keepNext')
            keepNext.set(qn('w:val'), '0')
            pPr.append(keepNext)
    
    # Helper function to ensure font size 12 for runs
    def set_font_size_12(run):
        """Set font size to 12 for a run"""
        run.font.size = Pt(12)
    
    # Helper function to set table border color
    def set_table_border_color(table):
        """Set border color to #256eb7 for all cells in a table"""
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '256EB7')
                    tc_borders.append(border)
                tc_pr.append(tc_borders)
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)  # Remove space after
    # Force zero spacing via XML
    pPr = p._element.get_or_add_pPr()
    for spacing_elem in pPr.xpath('.//w:spacing'):
        pPr.remove(spacing_elem)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    run = p.add_run('This Medication Assistance Plan is for:')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    # First box
    first_box = create_boxed_section()
    remove_all_spacing_from_cell(first_box)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Full Name:')
    run.bold = True
    set_font_size_12(run)
    # Add full name data
    full_name = f'{first_name} {surname}'.strip()
    if full_name:
        run = p.add_run(f' {full_name}')
        set_font_size_12(run)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(first_box)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Date of Birth:')
    set_font_size_12(run)
    # Add date of birth data
    if dob_str:
        run = p.add_run(f' {dob_str}')
        set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('NDIS Number:')
    set_font_size_12(run)
    # Add NDIS number data
    if ndis_number:
        p = add_paragraph_no_spacing(first_box)
        run = p.add_run(ndis_number)
        set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Medicare Number:')
    set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Alerts (Medic alert information etc.):')
    set_font_size_12(run)
    add_paragraph_no_spacing(first_box)  # Empty line
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('G.P. or Prescribing Doctor:')
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Name:')
    set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Contact Details:')
    set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Address:')
    set_font_size_12(run)
    add_paragraph_no_spacing(first_box)  # Empty line
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Pharmacist:')
    run.bold = True
    set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Name:')
    set_font_size_12(run)
    p = add_paragraph_no_spacing(first_box)
    run = p.add_run('Contact Details:')
    set_font_size_12(run)
    
    doc.add_paragraph()  # Empty line
    
    # Signature of Individual
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Signature of Individual (or person responsible for consent)')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    sig_box1 = create_boxed_section()
    remove_all_spacing_from_cell(sig_box1)
    p = add_paragraph_no_spacing(sig_box1)
    run = p.add_run('Date:')
    set_font_size_12(run)
    # Add four empty lines
    for _ in range(4):
        add_paragraph_no_spacing(sig_box1)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(sig_box1)
    
    doc.add_paragraph()  # Empty line
    
    # Plan Developed By
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Plan Developed By')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    plan_box = create_boxed_section()
    remove_all_spacing_from_cell(plan_box)
    p = add_paragraph_no_spacing(plan_box)
    run = p.add_run('Name of person responsible for developing the plan:')
    set_font_size_12(run)
    # Add key contact name
    if key_contact_name:
        p = add_paragraph_no_spacing(plan_box)
        run = p.add_run(key_contact_name)
        set_font_size_12(run)
    # Add four empty lines
    for _ in range(4):
        add_paragraph_no_spacing(plan_box)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(plan_box)
    
    doc.add_paragraph()  # Empty line
    
    # Signature of Authorised Medication Delegate
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Signature of Authorised Medication Delegate (Neighbourhood Care)')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    sig_box2 = create_boxed_section()
    remove_all_spacing_from_cell(sig_box2)
    p = add_paragraph_no_spacing(sig_box2)
    run = p.add_run('Date:')
    set_font_size_12(run)
    # Add four empty lines
    for _ in range(4):
        add_paragraph_no_spacing(sig_box2)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(sig_box2)
    
    doc.add_paragraph()  # Empty line
    
    # Reason For Plan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Reason For Plan')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    reason_box = create_boxed_section()
    remove_all_spacing_from_cell(reason_box)
    p = add_paragraph_no_spacing(reason_box)
    run = p.add_run('Please describe why a support plan is required.')
    set_font_size_12(run)
    # Add four empty lines
    for _ in range(4):
        add_paragraph_no_spacing(reason_box)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(reason_box)
    
    doc.add_paragraph()  # Empty line
    
    # Assistance Required
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Assistance Required')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    assist_box = create_boxed_section()
    remove_all_spacing_from_cell(assist_box)
    p = add_paragraph_no_spacing(assist_box)
    run = p.add_run('Describe the assistance required')
    set_font_size_12(run)
    
    # Add bullet points with data (in the order specified by user)
    assistance_items = []
    
    # 1. Communication assistance required
    if communication_assistance:
        assistance_items.append(f'• {communication_assistance}')
    
    # 2. Medication assistance (if Yes)
    if medication_assistance_needed and medication_assistance_needed.lower() in ['yes', 'y']:
        assistance_items.append('• Medication assistance')
    
    # 3. Equipment or assistive technologies used
    if equipment_assistive:
        assistance_items.append(f'• {equipment_assistive}')
    
    # 4. Assisted transfers
    if assisted_transfers:
        assistance_items.append(f'• {assisted_transfers}')
    
    # 5. Catheter management
    if catheter_management:
        assistance_items.append(f'• {catheter_management}')
    
    # 6. Wound care and/or pressure care
    if wound_pressure_care:
        assistance_items.append(f'• {wound_pressure_care}')
    if bowel_care:
        assistance_items.append(f'• {bowel_care}')
    if enteral_feeding:
        assistance_items.append(f'• {enteral_feeding}')
    if peg_feeding:
        assistance_items.append(f'• {peg_feeding}')
    if stoma_care:
        assistance_items.append(f'• {stoma_care}')
    if additional_care:
        assistance_items.append(f'• {additional_care}')
    if behaviour_support:
        assistance_items.append(f'• {behaviour_support}')
    
    # Add bullet points to the box
    for item in assistance_items:
        p = add_paragraph_no_spacing(assist_box)
        run = p.add_run(item)
        set_font_size_12(run)
    
    # Add four empty lines
    for _ in range(4):
        add_paragraph_no_spacing(assist_box)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(assist_box)
    
    doc.add_paragraph()  # Empty space before "Important things to remember"
    
    # Important things to remember
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Important things to remember')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    important_box = create_boxed_section()
    remove_all_spacing_from_cell(important_box)
    p = add_paragraph_no_spacing(important_box)
    run = p.add_run('Any additional plans relating to the person\'s medication should be listed here')
    set_font_size_12(run)
    # Add four empty lines
    for _ in range(4):
        add_paragraph_no_spacing(important_box)
    # Remove spacing again after adding content
    remove_all_spacing_from_cell(important_box)
    
    doc.add_paragraph()  # Empty line
    
    # Allergies & reactions
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Allergies & reactions')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    p = doc.add_paragraph('Any allergies (relating to medication) and potential reactions should be listed here.')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_font_size_12(run)
    
    # Allergies table
    allergies_table = doc.add_table(rows=6, cols=2)  # 1 header + 5 empty rows
    allergies_table.style = 'Table Grid'
    header_cells = allergies_table.rows[0].cells
    header_cells[0].paragraphs[0].add_run('What medications allergic to').bold = True
    header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_cells[1].paragraphs[0].add_run('Potential Reaction').bold = True
    header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Fill header cells with #e0f4ff
    for cell in header_cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E0F4FF')  # #e0f4ff
        shd.set(qn('w:val'), 'clear')
        tc_pr.append(shd)
    
    set_table_border_color(allergies_table)
    
    doc.add_paragraph()  # Empty line
    
    # Medications Prescribed & Potential Side Effects
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Medications prescribed & Potential Side Effects')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    # Side effects table
    side_effects_table = doc.add_table(rows=6, cols=2)  # 1 header + 5 empty rows
    side_effects_table.style = 'Table Grid'
    header_cells = side_effects_table.rows[0].cells
    header_cells[0].paragraphs[0].add_run('Medication').bold = True
    header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_cells[1].paragraphs[0].add_run('Side Effects').bold = True
    header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Fill header cells with #e0f4ff
    for cell in header_cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E0F4FF')  # #e0f4ff
        shd.set(qn('w:val'), 'clear')
        tc_pr.append(shd)
    
    set_table_border_color(side_effects_table)
    
    doc.add_paragraph()  # Empty line
    
    # Medication List - Prescribed
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Medication List - Prescribed - Update when changed by prescribing physician')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    # Prescribed medications table
    prescribed_table = doc.add_table(rows=1, cols=6)
    prescribed_table.style = 'Table Grid'
    
    # Set explicit column widths for the parent table
    # Total width should be ~7.5 inches (10800 twips) to fit on page
    # Column widths: Medication (1.3"), Dose (0.8"), When to take it (2.0"), How to take it (1.2"), Where (1.0"), Additional (0.7")
    col_widths = [1872, 1152, 60, 1728, 1440, 1008]  # In twips - "When to take it" is now minimal (60 twips)
    for i, width in enumerate(col_widths):
        header_cell = prescribed_table.rows[0].cells[i]
        tc_pr = header_cell._element.get_or_add_tcPr()
        # Remove existing width
        for width_elem in tc_pr.xpath('.//w:tcW'):
            tc_pr.remove(width_elem)
        # Set explicit width
        tc_width = OxmlElement('w:tcW')
        tc_width.set(qn('w:w'), str(width))
        tc_width.set(qn('w:type'), 'dxa')
        tc_pr.append(tc_width)
    
    header_cells = prescribed_table.rows[0].cells
    headers = ['Medication', 'Dose', 'When to take it', 'How to take it', 'Where it is kept', 'Additional details']
    for i, header_text in enumerate(headers):
        run = header_cells[i].paragraphs[0].add_run(header_text)
        # Bold only the specified headers
        if header_text in ['Medication', 'Dose', 'When to take it', 'How to take it', 'Where it is kept', 'Additional details']:
            run.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Fill header cells with #e0f4ff
        tc_pr = header_cells[i]._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E0F4FF')  # #e0f4ff
        shd.set(qn('w:val'), 'clear')
        tc_pr.append(shd)
    
    set_table_border_color(prescribed_table)
    
    # Add a data row with nested table in "When to take it" column
    data_row = prescribed_table.add_row()
    # Leave first 2 cells empty for now
    when_cell = data_row.cells[2]  # "When to take it" column
    when_cell.paragraphs[0].clear()
    
    # Remove all spacing and margins from the cell
    remove_all_spacing_from_cell(when_cell)
    
    # Set explicit width for the cell - just enough for the nested table
    # Nested table is only 50 twips wide, so set cell to match
    tc_pr = when_cell._element.get_or_add_tcPr()
    # Remove existing width
    for width_elem in tc_pr.xpath('.//w:tcW'):
        tc_pr.remove(width_elem)
    tc_width = OxmlElement('w:tcW')
    tc_width.set(qn('w:w'), '70')  # Just enough for nested table + minimal padding
    tc_width.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_width)
    
    # Create nested table for "When to take it"
    # Structure: Mini table
    # Row 1 (Heading 1): "Time" (spanning 2 cols) | "Day" (spanning 7 cols) = 9 columns total
    # Row 2 (Heading 2): AM | PM | S | M | T | W | T | F | S (all as horizontal headers)
    nested_table = when_cell.add_table(rows=2, cols=9)
    nested_table.style = 'Table Grid'
    
    # Column widths: AM (10), PM (10), S (6), M (6), T (6), W (6), T (6), F (6), S (6)
    # All single letters get the SAME width (6 twips), AM/PM get 10 twips each
    # Using VERY small widths to minimize spacing
    col_widths = [10, 10, 6, 6, 6, 6, 6, 6, 6]
    total_width = sum(col_widths)  # 62 twips
    
    # Set table properties FIRST - this is critical for nested tables
    tbl_pr = nested_table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        nested_table._element.insert(0, tbl_pr)
    
    # Remove any existing grid or width settings - be very aggressive
    for existing_grid in list(tbl_pr.xpath('.//w:tblGrid')):
        tbl_pr.remove(existing_grid)
    for existing_width in list(tbl_pr.xpath('.//w:tblW')):
        tbl_pr.remove(existing_width)
    for existing_layout in list(tbl_pr.xpath('.//w:tblLayout')):
        tbl_pr.remove(existing_layout)
    
    # Set table layout to fixed FIRST - CRITICAL for width enforcement
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)
    
    # Set table grid column widths - THIS IS THE PRIMARY SOURCE OF TRUTH FOR COLUMN WIDTHS
    # The grid defines the column structure
    tblGrid = OxmlElement('w:tblGrid')
    for width in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        gridCol.set(qn('w:type'), 'dxa')
        tblGrid.append(gridCol)
    tbl_pr.append(tblGrid)
    
    # Set table width to match grid - use exact sum
    tbl_width = OxmlElement('w:tblW')
    tbl_width.set(qn('w:w'), str(total_width))
    tbl_width.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_width)
    
    # NOW set widths on ALL cells in ALL rows to match the grid (after grid is set)
    # This must be done BEFORE any merging or content is added
    for row_idx, row in enumerate(nested_table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(col_widths):
                tc_pr = cell._element.get_or_add_tcPr()
                # Remove ALL existing width elements - be very aggressive
                for width_elem in list(tc_pr.xpath('.//w:tcW')):
                    tc_pr.remove(width_elem)
                # Set width to match grid - MUST use 'dxa' type
                tc_width = OxmlElement('w:tcW')
                tc_width.set(qn('w:w'), str(col_widths[col_idx]))
                tc_width.set(qn('w:type'), 'dxa')
                tc_pr.append(tc_width)
    
    # Header row 1: "Time" spanning 2 columns (first), "Day" spanning 7 columns (second) - no gap between them
    row0 = nested_table.rows[0]
    
    # Merge cells for "Time" first (cells 0-1)
    time_cell = row0.cells[0]
    if len(row0.cells) > 1:
        time_cell.merge(row0.cells[1])
    
    # Add "Time" text - increased font, not bold
    remove_all_spacing_from_nested_cell(time_cell)
    time_cell.paragraphs[0].clear()
    time_run = time_cell.paragraphs[0].add_run('Time')
    time_run.font.size = Pt(11)  # Increased font (11pt)
    time_run.bold = False  # Not bold
    time_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_all_spacing_from_nested_cell(time_cell)
    
    # Now merge "Day" cells - should be directly adjacent to Time with no gap
    # After merging Time, row0.cells[1] should be the first Day cell
    row0 = nested_table.rows[0]  # Get fresh reference
    if len(row0.cells) >= 2:
        day_cell = row0.cells[1]  # This should be directly after Time, no gap
        # Merge the remaining 6 cells into cell 1 to create Day spanning 7 columns
        for _ in range(6):
            if len(row0.cells) > 2:
                day_cell.merge(row0.cells[2])
        
        # Add "Day" text - increased font, not bold
        remove_all_spacing_from_nested_cell(day_cell)
        day_cell.paragraphs[0].clear()
        day_run = day_cell.paragraphs[0].add_run('Day')
        day_run.font.size = Pt(11)  # Increased font (11pt)
        day_run.bold = False  # Not bold
        day_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        remove_all_spacing_from_nested_cell(day_cell)
    
    # Second row (Heading 2): AM, PM (under Time), then S, M, T, W, T, F, S (under Day)
    nested_data = nested_table.rows[1].cells
    labels = ['AM', 'PM', 'S', 'M', 'T', 'W', 'T', 'F', 'S']  # Swapped order: Time fields first, then Day fields
    for i, label in enumerate(labels):
        # Remove spacing and margins from each cell using nested cell function
        remove_all_spacing_from_nested_cell(nested_data[i])
        nested_data[i].paragraphs[0].clear()
        label_run = nested_data[i].paragraphs[0].add_run(label)
        label_run.font.size = Pt(11)  # Increased font (11pt)
        label_run.bold = False  # Not bold
        nested_data[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Remove spacing again after adding content
        remove_all_spacing_from_nested_cell(nested_data[i])
        
        # CRITICAL: Re-apply width AFTER adding content to ensure it sticks
        # Use the same col_widths array defined earlier: [10, 10, 6, 6, 6, 6, 6, 6, 6]
        expected_widths = [10, 10, 6, 6, 6, 6, 6, 6, 6]
        if i < len(expected_widths):
            tc_pr = nested_data[i]._element.get_or_add_tcPr()
            # Remove ALL existing width elements
            for width_elem in tc_pr.xpath('.//w:tcW'):
                tc_pr.remove(width_elem)
            # Set explicit width
            tc_width = OxmlElement('w:tcW')
            tc_width.set(qn('w:w'), str(expected_widths[i]))
            tc_width.set(qn('w:type'), 'dxa')
            tc_pr.append(tc_width)
    
    set_table_border_color(nested_table)
    
    # FINAL: Force widths on row 1 cells one last time - use direct XML manipulation
    # Column widths: AM (10), PM (10), S (6), M (6), T (6), W (6), T (6), F (6), S (6)
    # ALL single letters MUST be the same width (6 twips)
    final_widths = [10, 10, 6, 6, 6, 6, 6, 6, 6]
    row1 = nested_table.rows[1]
    for i, cell in enumerate(row1.cells):
        if i < len(final_widths):
            # Get the cell's XML element directly
            tc = cell._element
            tc_pr = tc.get_or_add_tcPr()
            
            # Remove ALL existing width elements - be very aggressive
            for width_elem in list(tc_pr.xpath('.//w:tcW')):
                tc_pr.remove(width_elem)
            
            # Create and append width element
            tc_width = OxmlElement('w:tcW')
            tc_width.set(qn('w:w'), str(final_widths[i]))
            tc_width.set(qn('w:type'), 'dxa')
            tc_pr.append(tc_width)
            
            # Verify it was added
            verify_width = tc_pr.xpath('.//w:tcW')
            if not verify_width:
                # If it wasn't added, try inserting it directly
                tc_pr.insert(0, tc_width)
    
    # Set widths on row 0 merged cells AFTER merging
    row0 = nested_table.rows[0]
    if len(row0.cells) >= 2:
        # Time cell (first merged cell - spans columns 0-1)
        time_tc_pr = row0.cells[0]._element.get_or_add_tcPr()
        for width_elem in time_tc_pr.xpath('.//w:tcW'):
            time_tc_pr.remove(width_elem)
        time_width = OxmlElement('w:tcW')
        time_width.set(qn('w:w'), '20')  # 10 + 10 = 20 twips
        time_width.set(qn('w:type'), 'dxa')
        time_tc_pr.append(time_width)
        
        # Day cell (second merged cell - spans columns 2-8, directly adjacent to Time)
        day_tc_pr = row0.cells[1]._element.get_or_add_tcPr()
        for width_elem in day_tc_pr.xpath('.//w:tcW'):
            day_tc_pr.remove(width_elem)
        day_width = OxmlElement('w:tcW')
        day_width.set(qn('w:w'), '42')  # 6 * 7 = 42 twips
        day_width.set(qn('w:type'), 'dxa')
        day_tc_pr.append(day_width)
    
    doc.add_paragraph()  # Empty line
    
    # Medication List - As Needed (PRN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Medication List - As Needed (PRN)')
    run.bold = True
    run.font.color.rgb = text_color  # #007bc4
    set_font_size_12(run)
    
    # PRN medications table
    prn_table = doc.add_table(rows=6, cols=6)  # 1 header + 5 empty rows
    prn_table.style = 'Table Grid'
    header_cells = prn_table.rows[0].cells
    headers = ['Medication', 'What it is used for', 'Indications for use', 'How to take it/dose', 'Where it is kept', 'Additional details']
    for i, header_text in enumerate(headers):
        run = header_cells[i].paragraphs[0].add_run(header_text)
        # Bold only the specified headers
        if header_text in ['Medication', 'What it is used for', 'Indications for use', 'How to take it/dose', 'Where it is kept', 'Additional details']:
            run.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Fill header cells with #e0f4ff
        tc_pr = header_cells[i]._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E0F4FF')  # #e0f4ff
        shd.set(qn('w:val'), 'clear')
        tc_pr.append(shd)
    
    set_table_border_color(prn_table)
    
    # Final text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Observed Practice Checklist to be attached to this plan and records maintained by all parties involved in the medication assistance.')
    run.bold = True
    set_font_size_12(run)
    
    # Save document
    doc.save(output_path)
    print(f"Medication Assistance Plan DOCX created successfully: {output_path}")

if __name__ == "__main__":
    create_service_agreement()
